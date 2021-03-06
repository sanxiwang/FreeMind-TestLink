# -*- coding: utf-8 -*- 

import argparse
import logging.config
import sys
import os
from copy import deepcopy
from xml.etree import ElementTree as ET
import xml.dom.minidom as minidom
import xml.etree.cElementTree as xmlcET
from lxml import etree as lxmlET

import testlink
from xlrd import open_workbook
from xlwt import Formula, easyxf
from xlutils.copy import copy

from docx import Document
import pprint

PKG_PATH = './'

TC_ID = 0
TC_TITLE = 1
TC_REQ_LINKS = 2
REQ_LINK_SPEC = 0
REQ_LINK_ID = 1
REQ_LINK_TITLE = 2

REQ_ID = 0
REQ_TITLE = 1
REQ_DESC = 2
REQ_VER_TEAM = 3
REQ_COMMENT = 4
REQ_PHASE = 5

PREFIX_TITLE_SEP = '::'

''' The following functions (CDATA and _serialize_xml) is a workaround for using CDATA section with ElementTree
'''


def CDATA(text=None):
    element = ET.Element('![CDATA[')
    element.text = text
    return element


_original_serialize_xml = ET._serialize_xml


def _serialize_xml(write, elem, encoding, qnames, namespaces):
    if elem.tag == '![CDATA[':
        #write("<%s%s]]>%s" % (elem.tag, elem.text, elem.tail))
        write("<%s%s]]>" % (elem.tag, elem.text))
        return
    return _original_serialize_xml(
        write, elem, encoding, qnames, namespaces)


ET._serialize_xml = ET._serialize['xml'] = _serialize_xml


class FreeMind(object):
    ''' This is a class working with TestLink and various offline templates.
        Basically it includes the features of generating TDS, linking TDS with test cases and test plans.
        It also have some advance features including generating PMR, PFS and traceability from Excel template.
        Eventually it has the features of create test plan in test link and export various test reports from TestLink.
        Please check the related configuration file and work instructions for detailed information.
    '''

    def __init__(self, logger, cfg_file=None):
        self.logger = logger
        self.log_prefix = 'FreeMind:'
        self.fm_tree = None
        self.fm_file = None
        self.tc_tree = None
        self.tc_file = None
        self.node_found = False

        self.testlink_url = None
        self.testlink_devkey = None
        self.tls = None
        self.tc_prefix = None
        self.project_name = None
        self.pfs_prefix = None
        self.pmr_prefix = None
        self.tds_prefix = None
        self.test_plan = None
        self.requirements_url = None
        self.pmr_url = None
        self.pfs_url = None
        self.tds_url = None
        self.tp_url = None
        self.tc_url = None
        self.based_tc_url = None
        self.based_tp_url = None
        self.flashobject_swf = None
        self.flashobject_js = None
        self.html_template = None
        self.logger.info(self.log_prefix + \
                         "FreeMind-TestLink Tool 0.3 for Requirement Extract, Test Design and Test Management.")
        if cfg_file:
            # Parse the configuration file automatically if it's specified
            self.logger.info(self.log_prefix + \
                             "Parse the configuration file (%s)." % \
                             (cfg_file))
            self._parse_cfg_file(cfg_file)

    def _parse_cfg_file(self, cfg_file):
        cfg_tree = ET.parse(cfg_file)
        cfg_root = cfg_tree.getroot()

        # Firstly get all configurations from the default configuration file
        for item in cfg_root.iter():
            if item.tag == 'testlink':
                self.testlink_rpc_url = item.attrib['URL'].strip()
                self.testlink_url = '/'.join(self.testlink_rpc_url.split('/')[:4])
                self.testlink_devkey = item.attrib['DEV_KEY'].strip()
                os.environ['TESTLINK_API_PYTHON_SERVER_URL'] = self.testlink_rpc_url
                os.environ['TESTLINK_API_PYTHON_DEVKEY'] = self.testlink_devkey
            if item.tag == 'repository':
                self.repo_prefix = item.attrib['PREFIX'].strip()
                self.repo_name = item.attrib['NAME'].strip()
            if item.tag == 'project':
                self.project_name = item.attrib['NAME'].strip()
                self.pfs_prefix = item.attrib['PFS_PREFIX'].strip()
                if self.pfs_prefix == "":
                    self.pfs_prefix = self.project_name + '_'
                self.pmr_prefix = item.attrib['PMR_PREFIX'].strip()
                if self.pmr_prefix == "":
                    self.pmr_prefix = self.project_name + '_'
                self.tds_prefix = item.attrib['TDS_PREFIX'].strip()
                if self.tds_prefix == "":
                    self.tds_prefix = self.project_name + '_'

            if item.tag == 'file_location':
                file_location = item.attrib['URL'].strip()
            if item.tag == 'requirements_url':
                self.requirements_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'pmr_url':
                self.pmr_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'pfs_url':
                self.pfs_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'tds_url':
                self.tds_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'tc_url':
                self.tc_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'based_tc_url':
                self.based_tc_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'tp_url':
                self.tp_url = self._get_url(file_location, item.text.strip())
            if item.tag == 'based_tp_url':
                self.based_tp_url = self._get_url(file_location, item.text.strip())

            if item.tag == 'freemind':
                freemind = item.attrib['URL'].strip()
            if item.tag == 'flashobject_swf':
                self.flashobject_swf = freemind + item.text.strip()
            if item.tag == 'flashobject_js':
                self.flashobject_js = freemind + item.text.strip()
            if item.tag == 'html_template':
                self.html_template = freemind + item.text.strip()

                # Secondly perform all enabled actions.
        for action in cfg_root.iter('action'):
            if action.attrib['ENABLE'].strip() <> '1':
                continue
            action_name = action.attrib['NAME'].strip()
            self.logger.info(self.log_prefix + \
                             "Perform the enabled action (%s) specified in the configuration file (%s)." % \
                             (action_name, cfg_file))
            if action_name == 'Extract_Requirements':
                self.extract_requirements(self.requirements_url, action.attrib['TEMPLATE'].strip())
            if action_name == 'Extract_TestCases':
                self.extract_tc_from_file(self.tc_url, action.attrib['SHEET_NAME'].strip(), action.attrib['REVIEW_INFO'].strip())
            if action_name == 'Link_PFS_with_PMR':
                pass  #self.link_pfs_pmr(self.pmr_url, self.pfs_url)
            if action_name == 'Link_PFS_with_TCs':
                self.link_tc2pfs(action.attrib['TEAM'].strip())
            if action_name == 'Generate_TDS':
                self.gen_tds(self.tds_url, action.attrib['REMOVE_PREFIX'].strip())
            if action_name == 'Link_TDS_with_TCs':
                self.link_tc2tds(self.tds_url, self.tc_url)
            if action_name == 'Link_TDS_with_TCs-TPs':
                self.link_tp2tds_tc(self.tds_url, self.tc_url, action.attrib['FILTER'].strip())
            if action_name == 'Link_TDS_with_TCs-PFS':
                self.link_pfs2tds(self.tds_url, self.tc_url, self.pfs_url)
            if action_name == 'Link_TCs_with_TDS':
                self.link_tds2tc(self.tc_url, self.tds_url)
            if action_name == 'Create_Test_Plan':
                self.create_test_plan(self.tp_url, action.attrib['AUTO'].strip(), action.attrib['TEAM'].strip())
            if action_name == 'Generate_TCs_from_TDS':
                self.Generate_TCs_from_TDS(action.attrib['NODE_LIST'].strip(), action.attrib['TC_READY'].strip())
            if action_name == 'Check_PFS_Traceablity':
                self.chk_pfs_traceability(action.attrib['TEAM'].strip())
            if action_name == 'Generate_PFS_TC_Traceablity':
                self.gen_pfs_tc_traceability(action.attrib['TEAM'].strip())

        return 0

    def _get_url(self, file_location, file_name):
        ''' Combine the file location path with file names if they are sharing the same file location.
        '''
        res = None
        if os.path.exists(file_location):
            res = file_location + file_name
        else:
            res = file_name

        return res

    def parse_freemind(self, file_name):
        self.fm_tree = ET.parse(file_name)
        self.fm_file = file_name
        return 0

    def _gen_freemind(self):
        self.fm_tree.write(os.path.splitext(self.fm_file)[0] + "_New.mm")
        return 0

    def add_prefix(self, file_name):
        self.parse_freemind(file_name)
        self._add_node_prefix(self.fm_tree.getroot(), '0')
        self._gen_freemind()
        return 0

    def remove_prefix(self, file_name):
        self.parse_freemind(file_name)
        self._remove_node_prefix(self.fm_tree.getroot())
        self._gen_freemind()
        return 0

    def gen_tds(self, file_name, remove_prefix):
        tds_item_list = ['TDS', []]
        fm_tree = ET.parse(file_name)
        tds_root = fm_tree.getroot()
        #Firstly remove all prefix hence we will number them again.
        self._remove_node_prefix(tds_root)

        self.logger.info(self.log_prefix + \
                         "Read TDS file (%s) and get the information of last nodes which will be used to generate the xml file for importing to TestLink" % \
                         (file_name))
        self._get_tds_items(tds_root, '0', '', tds_item_list[1])

        filename = os.path.splitext(file_name)[0] + '.xml'
        title = os.path.splitext(os.path.split(file_name)[-1])[0]
        self._gen_req_xml([tds_item_list], title, filename, self.tds_prefix)

        self._update_pfs_node_format(tds_root)
        fm_tree.write(file_name)

        if remove_prefix == '1':
            self._remove_node_prefix(tds_root)
            fm_tree.write(file_name)

        return 0

    def _get_tds_items(self, node, num, desc, item_list):
        res = 0
        i = 0
        prefix = ''
        content = ''
        for child in node:
            if child.tag == 'node':
                if child.attrib.has_key('LINK') and child.attrib['LINK'].startswith(self.testlink_url):
                    continue
                i += 1
                prefix = str(num) + '.' + str(i)
                node_id = child.attrib['ID']
                content = desc + '|' + child.attrib['TEXT']
                # If this is the last TDS node
                if self._last_tds_node(child):
                    # Keep the TDS title as long as possible to about 100 characters (limitation in TestLink)
                    item_list.append(
                        [node_id, prefix[4:] + PREFIX_TITLE_SEP + '|'.join(content[-100:].split('|')[2:]), \
                         prefix[4:] + PREFIX_TITLE_SEP + '|'.join(content.split('|')[2:]), 'SIT'])
                    continue
                self._get_tds_items(child, prefix, content, item_list)

        return res

    def _gen_req_xml(self, item_list, doc_title, filename, prefix, relation_list=None):
        ''' item_list is a list like [GROUP_NAME, [ [REQ_ID, REQ_TITLE, REQ_DESC, REQ_VER_TEAM], ... ] ]
        '''
        res = 0

        self.logger.info(self.log_prefix + \
                         "Generating the xml file %s (Document Title: %s. Document ID Prefix: %s) for importing to TestLink." % \
                         (filename, doc_title, prefix))

        tds = ET.Element('requirement-specification')
        #title = ''.join(self.fm_file.split('.')[:-1])
        req_spec = ET.SubElement(tds, 'req_spec', {'title': doc_title, 'doc_id': doc_title})
        req_type = ET.SubElement(req_spec, 'type')
        req_type.append(CDATA(2))
        node_order = ET.SubElement(req_spec, 'node_order')
        node_order.append(CDATA(1))
        total_req = ET.SubElement(req_spec, 'total_req')
        total_req.append(CDATA(0))
        scope = ET.SubElement(req_spec, 'scope')
        scope.append(CDATA(''))

        #pprint.pprint(item_list)
        i = 0
        for group in item_list:
            for item in group[1]:
                i = i + 1
                requirement = ET.SubElement(req_spec, 'requirement')
                docid = ET.SubElement(requirement, 'docid')
                docid.append(CDATA(prefix + item[REQ_ID]))
                title = ET.SubElement(requirement, 'title')
                title.append(CDATA(item[REQ_TITLE]))
                node_order = ET.SubElement(requirement, 'node_order')
                node_order.append(CDATA(i))
                description = ET.SubElement(requirement, 'description')
                description.append(CDATA('<p>' + item[REQ_DESC].replace('\n', '</p><p>') + '</p>'))
                status = ET.SubElement(requirement, 'status')
                status.append(CDATA('V'))
                req_type = ET.SubElement(requirement, 'type')
                req_type.append(CDATA(2))
                expected_coverage = ET.SubElement(requirement, 'expected_coverage')
                expected_coverage.append(CDATA(1))
                custom_fields = ET.SubElement(requirement, 'custom_fields')
                custom_field = ET.SubElement(custom_fields, 'custom_field')
                name = ET.SubElement(custom_field, 'name')
                name.append(CDATA('HGI Req Verification Team'))
                value = ET.SubElement(custom_field, 'value')
                ver_team = item[REQ_VER_TEAM].split('\n')
                if len(ver_team) == 1:
                    ver_team = item[REQ_VER_TEAM].split(' ')
                if len(ver_team) == 1:
                    ver_team = item[REQ_VER_TEAM].split(',')
                if len(ver_team) == 1:
                    ver_team = item[REQ_VER_TEAM].split('|')
                if len(ver_team) == 1:
                    ver_team = item[REQ_VER_TEAM].split(';')
                ver_team = '|'.join(ver_team)
                value.append(CDATA(ver_team))

                if len(item) > REQ_COMMENT:
                    custom_field = ET.SubElement(custom_fields, 'custom_field')
                    name = ET.SubElement(custom_field, 'name')
                    name.append(CDATA('HGI Req Review Comments'))
                    value = ET.SubElement(custom_field, 'value')
                    value.append(CDATA(item[REQ_COMMENT]))
                    custom_field = ET.SubElement(custom_fields, 'custom_field')
                    name = ET.SubElement(custom_field, 'name')
                    name.append(CDATA('HGI Feature Phase'))
                    value = ET.SubElement(custom_field, 'value')
                    value.append(CDATA(item[REQ_PHASE]))

        if relation_list is not None:
            for relation_src in relation_list:
                for relation_dst in relation_src[1]:
                    relation = ET.SubElement(req_spec, 'relation')
                    source = ET.SubElement(relation, 'source')
                    source.text = relation_src[0]
                    destination = ET.SubElement(relation, 'destination')
                    destination.text = relation_dst
                    relation_type = ET.SubElement(relation, 'type')
                    relation_type.text = '1'

        rough_string = ET.tostring(tds, 'utf-8')
        #print rough_string
        reparsed = minidom.parseString(rough_string)
        f = open(filename, 'w')
        #reparsed.writexml(f, newl='\n', encoding='utf-8')
        reparsed.writexml(f, encoding='utf-8')
        f.close()

        self.logger.info(self.log_prefix + \
                         "xml file %s was generated successfully." % \
                         (filename))
        return res

    def link_pfs2tds(self, tds_url, tc_url, pfs_url):
        tc_req_list = []
        req_tc_list = []
        res = None

        res = self.link_tc2tds(tds_url, tc_url, tc_req_list, req_tc_list)

        #pprint.pprint(tc_req_list)
        pfs_file = os.path.splitext(self.pfs_url)[0] + '.mm'
        tds_tc_file = self.tds_url.replace('.mm', '[TDS-TC].mm')
        res = self._build_fm_traceability(tds_tc_file, pfs_file, tc_req_list,
                                          self.tds_url.replace('.mm', '[TDS-TC-PFS].mm'))

        return res

    def link_tc2tds(self, tds_file, tc_file, tc_req_list=None, req_tc_list=None):
        if tc_req_list == None:
            tc_req_list = []
        if req_tc_list == None:
            req_tc_list = []

        tc_fm_file = tc_file.replace('.xml', '.mm')
        res = self._read_tc_from_xml(tc_file, tc_fm_file, tc_req_list)
        res = self._reverse_links(tc_req_list, req_tc_list)
        #pprint.pprint(req_tc_list)

        fm_tree = ET.parse(tds_file)
        fm_root = fm_tree.getroot()
        #self._remove_node_prefix(fm_root)
        #self._add_node_prefix(fm_root, '0')
        #self._remove_link_node(fm_root)
        fm_tree.write(tds_file)
        #pprint.pprint(req_tc_list)
        res = self._build_fm_traceability(tds_file, tc_fm_file, req_tc_list, tds_file.replace('.mm', '[TDS-TC].mm'),
                                          True)

        return res


    def _link_tc_node(self, tc_id, tc_title, link_id, node):
        for child in node.iter('node'):
            # Make sure this is not the test case or requirement link node since only they are nodes with links
            if not child.attrib.has_key('LINK'):
                # If the TDS prefix matches with the link ID, then add this link to its sub-node
                if child.attrib['TEXT'].split(' ')[0] == link_id.split('_')[-1]:
                    link_text = 'HDVB-' + tc_id + ':' + tc_title
                    link_url = 'http://testlink.ea.mot.com/linkto.php?tprojectPrefix=HDVB&item=testcase&id=HDVB-' + tc_id
                    ET.SubElement(child, 'node', {'COLOR': '#990000', 'LINK': link_url, 'TEXT': link_text})
                    self.logger.info(self.log_prefix + \
                                     "Adding linkage sub-node (%s) to node (%s)" % \
                                     (link_text, child.attrib['TEXT']))
                    return 0

        self.logger.error(self.log_prefix + \
                          "Cannot find %s" % \
                          (tc_id))

        return None

    def _read_tc_from_xml(self, xml_file, fm_file, tc_req_list):
        tc_tree = xmlcET.parse(xml_file)
        tc_root = tc_tree.getroot()

        # Build the FreeMind for test case
        title = os.path.splitext(os.path.split(xml_file)[-1])[0]
        res = self._gen_tc_freemind(xml_file, title, fm_file)
        # Construct the traceability list between Test cases and Requirements/Test Design Specification  
        self.logger.info(self.log_prefix + \
                         "Getting traceability information from file %s" % \
                         (xml_file))
        prefix_list = [self.pmr_prefix, self.tds_prefix]
        #Could be multiple PFS prefix since some requirements will be reused between projects.
        prefix_list.extend(self.pfs_prefix.split('|'))
        for tc in tc_root.iter('testcase'):
            req_links = []
            tc_id = self.repo_prefix + '-' + str(tc.find('externalid').text)
            for req in tc.iter('requirement'):
                doc_id = req.find('doc_id').text

                for prefix in prefix_list:
                    # Check if this is a valid requirement/TDS for this project
                    if len(doc_id.split(prefix)) == 2:
                        req_links.append(doc_id.split(prefix)[1])
                        break
            # Please note the tc_id here is with the project prefix, and the req_id is without requirement prefix
            tc_req_list.append([tc_id, req_links])

        return res

    def _gen_tc_freemind(self, tc_file, title, output_file):
        ''' req_list is a list like [GROUP_NAME, [ [REQ_ID, REQ_TITLE, REQ_DESC, REQ_VER_TEAM], ... ] ]
            REQ_ID and REQ_TITLE will be combined as the node text and REQ_DESC will be displayed as comments
        '''
        tc_tree = xmlcET.parse(tc_file)
        tc_root = tc_tree.getroot()

        freemind = ET.Element('map', {'version': '1.0.1'})

        ET.SubElement(freemind, 'attribute_registry', {'SHOW_ATTRIBUTES': 'hide'})
        root_node = ET.SubElement(freemind, 'node', {'BACKGROUND_COLOR': '#0000ff', 'COLOR': '#000000', 'TEXT': title})
        ET.SubElement(root_node, 'font', {'NAME': 'SansSerif', 'SIZE': '20'})
        ET.SubElement(root_node, 'hook', {'NAME': 'accessories/plugins/AutomaticLayout.properties'})

        self._add_tc_details(tc_root, root_node)
        ET.ElementTree(freemind).write(output_file)
        self.logger.info(self.log_prefix + \
                         "Successfully generate test case FreeMind file %s" % \
                         (output_file))
        return 0

    def _add_tc_details(self, tc_root, fm_root):
        for child in tc_root:
            if child.tag == 'testsuite':
                #add a node in Freemind and call again.
                testsuite_node = ET.SubElement(fm_root, 'node',
                                               {'COLOR': '#990000', 'FOLDED': "true", 'TEXT': child.attrib['name']})
                ET.SubElement(testsuite_node, 'icon', {'BUILTIN': 'folder'})
                self._add_tc_details(child, testsuite_node)
                continue
            if child.tag == 'testcase':
                #add a node in Freemind
                valid_tc = False
                node_comment = ''
                node_text = child.attrib['name']
                expected_results = ''
                tc_id = ''
                regression_level = ''
                for item in child:
                    if item.tag == 'externalid':
                        tc_id = str(item.text)
                        node_text = self.repo_prefix + '-' + tc_id + PREFIX_TITLE_SEP + node_text
                    if item.tag == 'summary':
                        node_comment = '<p>Summary:</p>' + str(item.text) + '<p></p>'
                    if item.tag == 'preconditions':
                        node_comment = node_comment + '<p>Preconditions:</p>' + str(item.text) + '<p></p>'
                    if item.tag == 'steps':
                        node_comment = node_comment + '<p>Steps:</p>'
                        expected_results = '<p>Expected results:</p>'
                        for step in item.iter():
                            if step.tag == 'step_number':
                                node_comment = node_comment + '<p>' + step.text + '.'
                                expected_results = expected_results + '<p>' + step.text + '.'
                            if step.tag == 'actions':
                                node_comment = node_comment + str(step.text).replace('<p>', '', 1)
                            if step.tag == 'expected_results':
                                expected_results = expected_results + str(step.text).replace('<p>', '', 1)
                    if item.tag == 'custom_fields':
                        for custom_field in item:
                            if list(custom_field)[0].text == 'HGI Regression Level':
                                if list(custom_field)[1].text is None:
                                    regression_level = 0
                                else:
                                    regression_level = 6 - len(list(custom_field)[1].text.split('|'))
                                    #Enable this once the test case is linked with requirements
                                    #                    if item.tag == 'requirements':
                                    #                        for requirement in item:
                                    #                            doc_id = list(requirement)[1].text
                                    #                            for prefix in [self.pfs_prefix, self.pmr_prefix, self.tds_prefix]:
                                    #                                if len(doc_id.split(prefix)) == 2:
                                    #                                    valid_tc = True
                                    #                                    break
                                    #                if not valid_tc:
                                    #                    continue
                node_comment = node_comment + '<p></p>' + expected_results
                node_link = self.testlink_url + '/linkto.php?tprojectPrefix=' + self.repo_prefix + '&item=testcase&id=' + self.repo_prefix + '-' + tc_id
                tc_node = ET.SubElement(fm_root, 'node', {'COLOR': '#990000', 'LINK': node_link, 'TEXT': node_text})
                richcontent = ET.SubElement(tc_node, 'richcontent', {'TYPE': 'NOTE'})
                html = ET.SubElement(richcontent, 'html')
                ET.SubElement(richcontent, 'head')
                body = ET.SubElement(html, 'body')
                for section in node_comment.replace('</p>', '').split('<p>'):
                    comment = ET.SubElement(body, 'p')
                    comment.text = section

                ET.SubElement(tc_node, 'icon', {'BUILTIN': 'full-' + str(regression_level)})
        return 0

    def link_tds2tc(self, fm_file, tc_file):
        test_suite = []
        tc_tds_list = []
        tc_pfs_list = []

        #res = self._read_tc_from_xml(tc_file, test_suite, tc_tds_list, tc_pfs_list)
        #res = self._gen_tc_freemind(test_suite)

        # res = self._link_tds_tc(self.tds_url)
        # res = self._link_tds_pfs()
        # res = self._link_pfs_tc()

        self.fm_file = fm_file
        tds_title = os.path.split(os.path.splitext(self.fm_file)[0])[1]
        self.fm_tree = xmlcET.parse(fm_file)
        fm_root = self.fm_tree.getroot()

        #parser = lxmlET.XMLParser(False)
        self.tc_file = tc_file
        self.tc_tree = ET.parse(tc_file)
        tc_root = self.tc_tree.getroot()

        # Firstly put all test cases with requirements/TDS links into a list
        link_list = []
        self._get_link_node(fm_root, link_list)
        #pprint.pprint(link_list)

        #Secondly loop through all test cases and add the TDS linkage in
        for tc in tc_root.iter('testcase'):
            tc_name = tc.get('name')
            tc_id = tc.find('externalid').text
            for tds_link in link_list:
                if tc_id == tds_link[0].split('-')[-1]:
                    if 1:
                        tds_link_found = False
                        for req in tc.iter('requirement'):
                            if (req.find('req_spec_title').text == tds_title) and \
                                    (req.find('doc_id').text.split('_')[-1] == tds_link[3]):
                                tds_link_found = True
                                break
                        if not tds_link_found:
                            requirements = tc.find('requirements')
                            if requirements == None:
                                requirements = ET.SubElement(tc, 'requirements')
                            link_item = ET.SubElement(requirements, 'requirement')
                            req_spec_title = ET.SubElement(link_item, 'req_spec_title')
                            #req_spec_title.text = lxmlET.CDATA(tds_title)
                            req_spec_title.append(CDATA(tds_title))
                            doc_id = ET.SubElement(link_item, 'doc_id')
                            doc_id.append(CDATA(tds_link[2]))
                            #                            self.logger.info(self.log_prefix + \
                            #                                "Add TDS link (%s) in test case (%s:%s)" % \
                            #                                (tds_link[3], tc_id, tc_name)

                            #        filename = os.path.splitext(self.tc_file)[0] + "_New.xml"
                            #        rough_string = ET.tostring(tc_root, 'utf-8')
                            #        reparsed = minidom.parseString(rough_string)
                            #        f= open(filename, 'w')
                            #        reparsed.writexml(f, addindent='  ', newl='\n',encoding='utf-8')
                            #        f.close()

        self.tc_tree.write(os.path.splitext(self.tc_file)[0] + "_New.xml")
        return 0

    #    def create_test_plan(self, tp_url, based_tp_url, auto_sync, ver_team):
    #        ''' The inputs could be TDS aided test planning, Test Suites aided test planning or PFS aided test planning.
    #
    #        '''
    #        removed_tc_list = []
    #        tc_list = []
    #        new_tc_list = []
    #        fm_tree = ET.parse(tp_url)
    #        tp_root = fm_tree.getroot()
    #        based_tp_root = ET.parse(based_tp_url).getroot()
    #        #Firstly we need to go through the based test plan to see if there any node or test cases are removed from the based test plan.
    #        res = self._find_removed_tc(based_tp_root, tp_root, removed_tc_list)
    #        pprint.pprint(removed_tc_list)
    #        #Secondly we need to go through the new test plan and remove all the test cases based on information above, regression levels and verification teams.
    #        res = self._update_tp(tp_root, ver_team, removed_tc_list)
    #        #Remove the nodes without any test case
    #        res = self._remove_node_wo_tc(tp_root)
    #        #Get test case list
    #        res = self._get_fm_tc_list(tp_root, tc_list)
    #        #Remove duplicate test cases in list
    #        res = self._remove_duplicate(tc_list, new_tc_list)
    #        pprint.pprint(new_tc_list)
    #        #Generate the test plan for importing to TestLink
    #        # platform, name, id, version, order
    #
    #        #Create the test plan
    #        if auto_sync == '1':
    #            tp_name = os.path.split(os.path.splitext(tp_url)[0])[-1]
    #            res = self._create_test_plan_in_tl(tp_name, new_tc_list)
    #
    #        return res

    def gen_pfs_tc_traceability(self, ver_team):
        tc_req_list = []
        req_tc_list = []
        tc_fm_file = self.tc_url.replace('.xml', '.mm')
        res = self._read_tc_from_xml(self.tc_url, tc_fm_file, tc_req_list)
        res = self._reverse_links(tc_req_list, req_tc_list)
        #pprint.pprint(req_tc_list)
        self._update_pfs_with_tc_traceability(self.requirements_url, req_tc_list)

    def _update_pfs_with_tc_traceability(self, pfs_url, req_tc_list):
        self.logger.info(self.log_prefix + \
                         "Reading requirement file (%s) and updating traceability. This is going to take a while..." % \
                         (pfs_url))
        src_wb = open_workbook(pfs_url, formatting_info=True)
        for index, s in enumerate(src_wb.sheets()):
            if s.name.lower().count('specification') > 0:
                src_req_sheet = s
                break
        dst_wb = copy(src_wb)
        dst_req_sheet = dst_wb.get_sheet(index)
        plain = easyxf('')

        pfs_index_col = 0
        pfs_tc_col = 0
        col_defined = False
        for i, cell in enumerate(src_req_sheet.col(0)):
            if not col_defined:
                for j in range(0, src_req_sheet.ncols):
                    cell_text = str(src_req_sheet.cell_value(i, j)).strip()
                    if cell_text.lower() == 'index':
                        pfs_index_col = j
                    if cell_text.lower() == 'si&t':
                        ver_sit_col = j
                        col_defined = True
                        coverage_formula = 'COUNTA(' + unichr(ord('A')+pfs_tc_col) + str(i+2) + ':' + \
                                           unichr(ord('A')+pfs_tc_col) + str(src_req_sheet.nrows+1) + ')/COUNTA('+ \
                                           unichr(ord('A')+ver_sit_col) + str(i+2) + ':' + unichr(ord('A')+ver_sit_col) + \
                                           str(src_req_sheet.nrows+1) + ')'
                        dst_req_sheet.write(i, pfs_tc_col, Formula(coverage_formula), plain)
                        #print i+1,unichr(ord('A')+pfs_tc_col), coverage_formula
                    if cell_text.lower() == 'si&t coverage':
                        pfs_tc_col = j
                continue

            pfs_index = str(src_req_sheet.cell_value(i, pfs_index_col)).strip()
            if pfs_index == '':
                continue
            for req_item in req_tc_list:
                if req_item[0] == pfs_index:
                    pfs_tc_traceability = ', '.join(req_item[1])
                    dst_req_sheet.write(i, pfs_tc_col, pfs_tc_traceability, plain)

        output_file_name = pfs_url.replace(os.path.splitext(pfs_url)[-1], '[PFS-TC].xls')
        dst_wb.save(output_file_name)
        self.logger.info(self.log_prefix + \
                         "Successfully generated PFS-TC traceaility file (%s)" % \
                         (output_file_name))

    def chk_pfs_traceability(self, ver_team):
        """
        This function will check the traceability between PFS and TDS items. Only PFS applied to specified verification
        team will be checked and marked.
        """
        tc_pfs_dict = {}
        pfs_tc_dict = {}
        pfs_tree = lxmlET.parse(self.pfs_url.replace('.xml', '.mm'))
        pfs_root = pfs_tree.getroot()

        tds_tree = lxmlET.parse(self.tds_url)
        tds_root = tds_tree.getroot()
        res = self._get_tc_pfs_traceability(tds_root, tc_pfs_dict)
        self._reverse_dict(tc_pfs_dict, pfs_tc_dict)

        ver_team = ver_team.split('|')
        ver_team_list = [item.strip() for item in ver_team]
        for pfs_node in pfs_root.iter('node'):
            if pfs_node.attrib.has_key('LINK') and pfs_node.attrib['LINK'].startswith(self.testlink_url) and \
                            pfs_node.attrib['LINK'].count('req&id') > 0:
                pfs_ver_team = pfs_node.attrib['TEXT'].split(PREFIX_TITLE_SEP)[1]
                pfs_ver_team = pfs_ver_team.split('|')
                pfs_id = pfs_node.attrib['LINK'].split('=')[-1]
                for ver_team in ver_team_list:
                    if ver_team in pfs_ver_team:
                        if not pfs_tc_dict.has_key(pfs_id):
                            self.logger.error(self.log_prefix + \
                                              "PFS item (%s) with verification team (%s) doesn't have a traceable TDS item. Highlights it with red backgroud color" % \
                                              (pfs_id, pfs_ver_team))
                            pfs_node.set('BACKGROUND_COLOR', '#ff0000')
                        else:
                            self.logger.info(self.log_prefix + \
                                             "PFS item (%s) with verification team (%s) has %d TDS items traced." % \
                                             (pfs_id, pfs_ver_team, len(pfs_tc_dict[pfs_id])))

        pfs_tree.write(self.pfs_url.replace('.xml', '[PFS-TDS].mm'))

    def _reverse_dict(self, src_dict, dst_dict):
        """
        This function will reverse the traceability dictionary. The source dictionary has format like this:
        {SRC_ID1: [DST_ID1, DST_ID2], SRC_ID2:[DST_ID1, DST_ID2]}
        """
        for id, value_list in src_dict.iteritems():
            for value in value_list:
                if dst_dict.has_key(value):
                    dst_dict[value].append(id)
                else:
                    dst_dict[value] = [id]


    def Generate_TCs_from_TDS(self, node_list, tc_ready):
        """
        It will generate test cases from the last tds item node. It would be empty test case in testlink.
        However, it will create the traceability between this test case and PFS/TDS automatically in testlink.
        The generated xml file need to be imported into testlink manually.
        """
        tc_tds_dict = {}
        tc_pfs_dict = {}

        fm_tree = lxmlET.parse(self.tds_url)
        tds_root = fm_tree.getroot()
        node_list = node_list.split('|')
        node_list = [item.strip() for item in node_list]
        # Create traceability dictionary for last TDS nodes. (Including traceability to both PFS and TDS)
        res = self._get_tc_tds_traceability(tds_root, tc_tds_dict)
        res = self._get_tc_pfs_traceability(tds_root, tc_pfs_dict)
        #pprint.pprint(tc_pfs_dict)
        # Generate test cases automatically with traceability
        tc_root = lxmlET.Element('testsuite', {'name': ''})
        lxmlET.SubElement(tc_root, 'node_order').text = lxmlET.CDATA('')
        lxmlET.SubElement(tc_root, 'details').text = lxmlET.CDATA('')
        res = self._gen_tc_xml_from_tds(tc_root, tds_root, tc_tds_dict, tc_pfs_dict, node_list, tc_ready)
        f = open(self.tc_url, 'w')
        f.write(lxmlET.tostring(tc_root, xml_declaration=True, encoding='UTF-8', pretty_print=True))
        f.close
        self.logger.info(self.log_prefix + \
                         "Successfully generated the test cases xml file (%s)." % \
                         (self.tc_url))

        res = self._update_pfs_node_format(tds_root)
        fm_tree.write(self.tds_url)
        self.logger.info(self.log_prefix + \
                         "Updated PFS nodes in  TDS document (%s)." % \
                         (self.tds_url))

    def _update_pfs_node_format(self, tds_root):
        for tds_item in tds_root.iter('node'):
            if tds_item.attrib.has_key('LINK') and tds_item.attrib['LINK'].startswith(self.testlink_url) and \
                            tds_item.attrib['LINK'].count('req&id') > 0:
                tds_item.attrib['TEXT'] = tds_item.attrib['TEXT'].split(PREFIX_TITLE_SEP)[0]
                tds_item.attrib['BACKGROUND_COLOR'] = '#ffffff'
                tds_item.attrib['COLOR'] = '#00b439'
                tds_item.attrib['STYLE'] = 'bubble'
                font = tds_item.find('font')
                if font is not None:
                    font.attrib['NAME'] = 'SansSerif'
                    font.attrib['SIZE'] = '8'
                else:
                    lxmlET.SubElement(tds_item, 'font', {'NAME': 'SansSerif', 'SIZE': '8'})
                edge = tds_item.find('edge')
                if edge is not None:
                    edge.attrib['STYLE'] = 'bezier'
                    edge.attrib['WIDTH'] = 'thin'
                else:
                    lxmlET.SubElement(tds_item, 'edge', {'STYLE': 'bezier', 'WIDTH': 'thin'})


    def _gen_tc_xml_from_tds(self, ts_node, root_node, tc_tds_dict, tc_pfs_dict, node_list, tc_ready):
        existing_tc_list = []
        if node_list == ['']:
            self.logger.info(self.log_prefix + \
                             "Generating test cases xml file for all TDS nodes.")
            self._gen_tc_xml_from_tds_node(ts_node, root_node, tc_tds_dict, tc_pfs_dict, existing_tc_list, tc_ready)
            return
        for tds_item in root_node.iter('node'):
            if tds_item.attrib['ID'] in node_list:
                self.logger.info(self.log_prefix + \
                                 "Generating test cases xml file for TDS node (%s)." % \
                                 (tds_item.attrib['ID']))
                child_ts_node = lxmlET.SubElement(ts_node, 'testsuite', {'name': tds_item.attrib['TEXT'].strip()})
                lxmlET.SubElement(child_ts_node, 'node_order').text = lxmlET.CDATA('')
                lxmlET.SubElement(child_ts_node, 'details').text = lxmlET.CDATA('')
                self._gen_tc_xml_from_tds_node(child_ts_node, tds_item, tc_tds_dict, tc_pfs_dict, existing_tc_list,
                                               tc_ready)

    def _gen_tc_xml_from_tds_node(self, ts_node, root_node, tc_tds_dict, tc_pfs_dict, existing_tc_list, tc_ready):
        ts_node_order = -1
        tc_node_order = -1
        for tds_item in root_node.findall('node'):
            if tds_item.attrib.has_key('LINK') and tds_item.attrib['LINK'].startswith(self.testlink_url):
                continue
            is_testsuite = False
            for item_icon in tds_item.findall('icon'):
                if item_icon.attrib['BUILTIN'] == 'folder':
                    ts_node_order += 1
                    child_ts_node = lxmlET.SubElement(ts_node, 'testsuite', {'name': tds_item.attrib['TEXT'].strip()})
                    lxmlET.SubElement(child_ts_node, 'node_order').text = lxmlET.CDATA(str(ts_node_order))
                    lxmlET.SubElement(child_ts_node, 'details').text = lxmlET.CDATA('')
                    #ts_node = child_ts_node
                    is_testsuite = True
                    break
            if is_testsuite:
                self._gen_tc_xml_from_tds_node(child_ts_node, tds_item, tc_tds_dict, tc_pfs_dict, existing_tc_list,
                                               tc_ready)
                continue
            if not self._last_tds_node(tds_item):
                self._gen_tc_xml_from_tds_node(ts_node, tds_item, tc_tds_dict, tc_pfs_dict, existing_tc_list, tc_ready)
                continue
            # This must be the last TDS node
            tc_list = []
            res = self._get_linked_tc(tds_item, tc_list)
            if not tc_list:
                # There is no linked test case nodes (which mainly used for reusing test cases between projects)
                if tds_item.attrib['ID'].strip() in existing_tc_list:
                    continue
                existing_tc_list.append(tds_item.attrib['ID'].strip())
                tc_node_order += 1
                if tc_ready:
                    # Test cases for some of the nodes are ready in a xml file (for instance, tester has created
                    # test cases in an excel file or in testlink). We can use the test case name to associate them.
                    tc_node = self._get_tc_node_from_xml_by_name(self.based_tc_url, tds_item.attrib['TEXT'].strip())
                    if tc_node is None:
                        # If we don't have a test case for this TDS node, create a dummy test case.
                        res = self._add_dummy_testcase(ts_node, tds_item, tc_tds_dict, tc_pfs_dict, tc_node_order)
                        continue
                    res = self._update_tc_node(tc_node, tc_node_order, tds_item, tc_tds_dict, tc_pfs_dict)
                    ts_node.append(tc_node)
                else:
                    # If this node doesn't have a test case associated, create a new dummy test case with traceability.
                    res = self._add_dummy_testcase(ts_node, tds_item, tc_tds_dict, tc_pfs_dict, tc_node_order)
                continue
            # If this node already have test cases associated, update its traceability if necessary.
            # Get the test case from original xml file and copy it into the new xml file
            for tc_id in tc_list:
                if tc_id in existing_tc_list:
                    continue
                existing_tc_list.append(tc_id)
                tc_node_order += 1
                tc_node = self._get_tc_node_from_xml_by_id(self.based_tc_url, tc_id)
                if tc_node is None:
                    return
                res = self._update_tc_node(tc_node, tc_node_order, tds_item, tc_tds_dict, tc_pfs_dict, tc_id)
                ts_node.append(tc_node)

    def _get_tc_node_from_xml_by_id(self, xml_file, tc_id):
        parser = lxmlET.XMLParser(strip_cdata=False)
        tc_root = lxmlET.parse(xml_file, parser)
        for tc_node in tc_root.iter('testcase'):
            if tc_node.find('externalid').text == tc_id.split('-')[-1]:
                return tc_node
        self.logger.warning(self.log_prefix + \
                         "Test case (%s) can not be found in file (%s)." % \
                         (tc_id, xml_file))
        return None

    def _get_tc_node_from_xml_by_name(self, xml_file, tc_name):
        parser = lxmlET.XMLParser(strip_cdata=False)
        tc_root = lxmlET.parse(xml_file, parser)
        for tc_node in tc_root.iter('testcase'):
            if tc_node.attrib['name'].strip() == tc_name:
            #if tc_node.attrib['name'].strip().count(tc_name) > 0: # For Li Tong Only
            # if tc_name.count(tc_node.attrib['name'].strip()) > 0:# For Teng Chong Only
            #     tc_node.attrib['name'] = tc_name # For Teng Chong Only
                return tc_node
        self.logger.warning(self.log_prefix + \
                         "Test case (%s) can not be found in file (%s)." % \
                         (tc_name, xml_file))
        return None

    def _update_tc_node(self, tc_node, tc_node_order, tds_item, tc_tds_dict, tc_pfs_dict, tc_id=None):
        """
        Update traceability in this test case node
        TODO: If this test case is copied from another project (Can be known from tc_id),
        need to update internalid, node_order, externalid, version as well
        """
        tc_node.find('node_order').text = lxmlET.CDATA(str(tc_node_order))
        requirements = tc_node.find('requirements')
        if requirements is not None:
            for requirement in requirements.findall('requirement'):
                requirements.remove(requirement)
        else:
            requirements = lxmlET.SubElement(tc_node, 'requirements')
        requirement = lxmlET.SubElement(requirements, 'requirement')
        lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
            os.path.splitext(os.path.split(self.tds_url)[-1])[0])
        lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(tc_tds_dict[tds_item.attrib['ID']][0])
        if not tc_pfs_dict.has_key(tds_item.attrib['ID']):
            return
        for pfs_id in tc_pfs_dict[tds_item.attrib['ID']]:
            requirement = lxmlET.SubElement(requirements, 'requirement')
            lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
                os.path.splitext(os.path.split(self.pfs_url)[-1])[0])
            lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(pfs_id)

    def _add_dummy_testcase(self, ts_node, tds_item, tc_tds_dict, tc_pfs_dict, tc_node_order):
        if not tds_item.attrib.has_key('TEXT'):
            self.logger.error(self.log_prefix + \
                             "Please check node (%s) since it may use a long name. Please convert it to plain text via FreeMind Menu Format=>Use Plaine Text." % \
                             (tds_item.attrib['ID'].strip()))
            exit(-1)
        testcase = lxmlET.SubElement(ts_node, 'testcase', {'name': tds_item.attrib['TEXT'].strip()})
        lxmlET.SubElement(testcase, 'node_order').text = lxmlET.CDATA(str(tc_node_order))
        lxmlET.SubElement(testcase, 'externalid').text = lxmlET.CDATA('')
        lxmlET.SubElement(testcase, 'version').text = lxmlET.CDATA('1')
        lxmlET.SubElement(testcase, 'summary').text = lxmlET.CDATA('')
        lxmlET.SubElement(testcase, 'preconditions').text = lxmlET.CDATA('')
        lxmlET.SubElement(testcase, 'execution_type').text = lxmlET.CDATA('1')
        lxmlET.SubElement(testcase, 'importance').text = lxmlET.CDATA('3')

        steps = lxmlET.SubElement(testcase, 'steps')
        # step = lxmlET.SubElement(steps, 'step')
        # lxmlET.SubElement(step, 'step_number').text = lxmlET.CDATA('1')
        # lxmlET.SubElement(step, 'actions').text = lxmlET.CDATA('')
        # lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA('')
        # lxmlET.SubElement(step, 'execution_type').text = lxmlET.CDATA('1')
        #
        # custom_fields = lxmlET.SubElement(testcase, 'custom_fields')
        # custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        # lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Regression Level')
        # lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('')
        # custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        # lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Test Team')
        # lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('SIT')
        # custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        # lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed')
        # lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('')
        # custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        # lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed Version')
        # lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('')
        # custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        # lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Review Info')
        # lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('')

        requirements = lxmlET.SubElement(testcase, 'requirements')
        requirement = lxmlET.SubElement(requirements, 'requirement')
        lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
            os.path.splitext(os.path.split(self.tds_url)[-1])[0])
        lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(tc_tds_dict[tds_item.attrib['ID']][0])
        if not tc_pfs_dict.has_key(tds_item.attrib['ID']):
            return
        for pfs_id in tc_pfs_dict[tds_item.attrib['ID']]:
            requirement = lxmlET.SubElement(requirements, 'requirement')
            lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
                os.path.splitext(os.path.split(self.pfs_url)[-1])[0])
            lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(pfs_id)

    def _add_codecs_testcase(self, ts_node, tds_item, tc_tds_dict, tc_pfs_dict, tc_node_order):
        if not tds_item.attrib.has_key('TEXT'):
            self.logger.error(self.log_prefix + \
                             "Please check node (%s) since it may use a long name. Please convert it to plain text via FreeMind Menu Format=>Use Plaine Text." % \
                             (tds_item.attrib['ID'].strip()))
            exit(-1)
        testcase = lxmlET.SubElement(ts_node, 'testcase', {'name': tds_item.attrib['TEXT'].strip()})
        lxmlET.SubElement(testcase, 'node_order').text = lxmlET.CDATA(str(tc_node_order))
        lxmlET.SubElement(testcase, 'externalid').text = lxmlET.CDATA('')
        lxmlET.SubElement(testcase, 'version').text = lxmlET.CDATA('1')
        # Verify the audio format of MPEG-4 AAC-HE  [VBR] Bitrate:100 kbps is decoded and streamed from the all applied audio outputs
        lxmlET.SubElement(testcase, 'summary').text = lxmlET.CDATA('Verify the audio format of ' + tds_item.attrib[
            'TEXT'].strip() + ' is decoded and streamed from the all applied audio outputs.')
        # lxmlET.SubElement(testcase, 'summary').text = lxmlET.CDATA('Verify the video format of ' + tds_item.attrib[
        #     'TEXT'].strip() + ' is displayed without visible artifacts, tiling or distortion.')
        lxmlET.SubElement(testcase, 'preconditions').text = lxmlET.CDATA('')
        lxmlET.SubElement(testcase, 'execution_type').text = lxmlET.CDATA('1')
        lxmlET.SubElement(testcase, 'importance').text = lxmlET.CDATA('1')

        steps = lxmlET.SubElement(testcase, 'steps')
        step = lxmlET.SubElement(steps, 'step')
        lxmlET.SubElement(step, 'step_number').text = lxmlET.CDATA('1')
        #Play the stream format of MPEG-4 AAC-HE  [VBR] Bitrate:100 kbps.
        lxmlET.SubElement(step, 'actions').text = lxmlET.CDATA(
            'Play the stream with format of ' + tds_item.attrib['TEXT'].strip() + '.')
        #AAC-HE format is decoded and streamed from the all applied audio outputs
        lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA(
            'Audio is decoded and streamed from the all applied audio outputs.')
        # lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA(
        #     'Video is displayed without visible artifacts, tiling or distortion.')
        lxmlET.SubElement(step, 'execution_type').text = lxmlET.CDATA('1')

        custom_fields = lxmlET.SubElement(testcase, 'custom_fields')
        custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Regression Level')
        lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('5 - First Time Run')
        custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Test Team')
        lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('SIT')
        custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed')
        lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('Yes')
        custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed Version')
        lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('1')
        custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
        lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Review Info')
        lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('Reviewed by Anderson Wang on 2014/4/25.')

        requirements = lxmlET.SubElement(testcase, 'requirements')
        requirement = lxmlET.SubElement(requirements, 'requirement')
        lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
            os.path.splitext(os.path.split(self.tds_url)[-1])[0])
        lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(tc_tds_dict[tds_item.attrib['ID']][0])
        if not tc_pfs_dict.has_key(tds_item.attrib['ID']):
            return
        for pfs_id in tc_pfs_dict[tds_item.attrib['ID']]:
            requirement = lxmlET.SubElement(requirements, 'requirement')
            lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
                os.path.splitext(os.path.split(self.pfs_url)[-1])[0])
            lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(pfs_id)

    def _get_tc_pfs_traceability(self, root_node, tc_pfs_dict):
        self.logger.info(self.log_prefix + \
                         "Getting traceability between PFS and TDS items.")
        for tds_node in root_node.iter('node'):
            if tds_node.attrib.has_key('LINK'):
                if tds_node.attrib['LINK'].startswith(self.testlink_url) and tds_node.attrib['LINK'].count(
                        'req&id') > 0:
                    # This is a PFS node, so all valid TDS items under the parent node of this node will have this PFS ID as traceability.
                    self._add_tc_pfs_traceability(tds_node.getparent(), tc_pfs_dict,
                                                  tds_node.attrib['LINK'].split('=')[-1])

    def _add_tc_pfs_traceability(self, root_node, tc_pfs_dict, pfs_id):
        for tds_item in root_node.iter('node'):
            if not self._last_tds_node(tds_item):
                continue
            if tc_pfs_dict.has_key(tds_item.attrib['ID']):
                duplicated_pfs = False
                for orig_pfs_id in tc_pfs_dict[tds_item.attrib['ID']]:
                    if orig_pfs_id == pfs_id:
                        duplicated_pfs = True
                        self.logger.warning(self.log_prefix + \
                                            "Duplicated PFS item (%s) found for TDS node (%s:%s)" % \
                                            (pfs_id, tds_item.attrib['ID'], tds_item.attrib['TEXT']))
                        break
                if not duplicated_pfs:
                    tc_pfs_dict[tds_item.attrib['ID']].append(pfs_id)
            else:
                tc_pfs_dict[tds_item.attrib['ID']] = [pfs_id]

    def _get_tc_tds_traceability(self, root_node, tc_tds_dict):
        self.logger.info(self.log_prefix + \
                         "Getting traceability between test cases and TDS items.")
        for tds_item in root_node.iter('node'):
            if not self._last_tds_node(tds_item):
                continue
            # If this is the last node and a node with only PFS items (we called 'valid tds item'), then this is a valid node that will be imported into testlink for traceability.
            if not tc_tds_dict.has_key(tds_item.attrib['ID']):
                tc_tds_dict[tds_item.attrib['ID']] = [self.tds_prefix + tds_item.attrib['ID']]
                #print tds_item.attrib['TEXT']
            else:
                self.logger.error(self.log_prefix + \
                                  "Duplicated TDS item (%s) found. Please check your FreeMind file in text mode." % \
                                  (tds_item.attrib['ID']))

    def _last_tds_node(self, node):
        if node.attrib.has_key('LINK') and node.attrib['LINK'].startswith(self.testlink_url):
            # This maybe a PFS/TC item, we need ignore it
            return False
        for child in node.findall('node'):
            if not (child.attrib.has_key('LINK') and child.attrib['LINK'].startswith(self.testlink_url)):
                return False
        return True

    def _get_linked_tc(self, tds_item, tc_list):
        for child in tds_item.findall('node'):
            if child.attrib.has_key('LINK') and child.attrib['LINK'].startswith(self.testlink_url) and child.attrib[
                'LINK'].count('testcase&id') > 0:
                tc_list.append(child.attrib['LINK'].split('=')[-1].strip())

    def create_test_plan(self, tp_url, auto_sync, ver_team):
        ''' The inputs could be TDS aided test planning, Test Suites aided test planning or PFS aided test planning.                        
        '''
        removed_tc_list = []
        kept_tc_list = []
        tc_list = []
        new_tc_list = []
        fm_tree = ET.parse(tp_url)
        tp_root = fm_tree.getroot()

        #Firstly we need to go through the test plan to see if there any test case is removed or there are any test cases need to be kept.
        res = self._find_removed_kept_tc(tp_root, removed_tc_list, kept_tc_list)
        self.logger.info(self.log_prefix + \
                         "Test cases marked with remove icon are (%s)." % \
                         (removed_tc_list))
        self.logger.info(self.log_prefix + \
                         "Test cases marked with must-keep icon are (%s)." % \
                         (kept_tc_list))
        #Secondly we need to get all test cases based on information above, regression levels and verification teams.
        res = self._get_tc_list(tp_root, removed_tc_list, kept_tc_list, tc_list, ver_team)
        res = self._remove_duplicate(tc_list, new_tc_list)
        self.logger.info(self.log_prefix + \
                         "Test cases planned in this test cycle are (%s)." % \
                         (new_tc_list))

        #Update Test Plan
        res = self._update_fm_tp(tp_root, new_tc_list)
        fm_tree.write(tp_url)
        self.logger.info(self.log_prefix + \
                         "The original test plan file (%s) is updated." % \
                         (tp_url))

        #Generate the test plan for importing to TestLink
        # platform, name, id, version, order

        #Create the test plan
        if auto_sync == '1':
            tp_name = os.path.split(os.path.splitext(tp_url)[0])[-1]
            res = self._create_test_plan_in_tl(tp_name, new_tc_list)

        return res

    def _create_test_plan_in_tl(self, tp_name, tc_list):
        ''' Establish a connection with TestLink and then create a new test plan.
            Get the latest version the assigned test cases and then add them into the test plan.
            It could be very slow depending on the link and xmlrpc.
        '''
        self.logger.info(self.log_prefix + \
                         "Test plan (%s) will be created and updated in TestLink. This is going to take a while. Please wait..." % \
                         (tp_name))
        self.tls = testlink.TestLinkHelper().connect(testlink.TestlinkAPIClient)
        prj = self.tls.getTestProjectByName(self.repo_name)
        prj_id = prj['id']
        tp = self.tls.createTestPlan(tp_name, self.repo_name)
        tp_id = tp[0]['id']
        #tp_id = self.tls.getTestPlanByName(self.repo_name, tp_name)[0]['id']
        for tc_id in tc_list:
            tc_version = self.tls.getTestCase(None, testcaseexternalid=tc_id)[0]['version']
            self.tls.addTestCaseToTestPlan(prj_id, tp_id, tc_id, int(tc_version))

        self.logger.info(self.log_prefix + \
                         "Test plan (%s) is created and updated successfully." % \
                         (tp_name))

    def link_tp2tds_tc(self, tds_url, tc_url, name_filter):
        tc_list = []
        res = self._get_test_plan_info(name_filter, tc_list)
        #pprint.pprint(tc_list) 
        # Link TDS_TC file with Test Plan and Execution status
        #res = self.link_tc2tds(self.tds_url, self.tc_url)
        res = self._link_tp2fm(tds_url.replace('.mm', '[TDS-TC].mm'), tc_list)

    def _link_tp2fm(self, fm_file, tc_list):
        tp_list = []
        fm_tree = ET.parse(fm_file)
        root_node = fm_tree.getroot()
        for child in root_node.iter('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case            
            if (tc_id.count(self.repo_prefix) == 1):
                tp_list = []
                for tc in tc_list:
                    if tc[0] == tc_id:
                        tp_list = tc[1]
                        break
                #print tp_list
                for tp in tp_list:
                    tp_name = tp[0]
                    tp_sts = tp[1]
                    tp_node = ET.SubElement(child, 'node', {'TEXT': tp_name})
                    if tp_sts == 'p':
                        ET.SubElement(tp_node, 'icon', {'BUILTIN': 'go'})
                    if tp_sts == 'f':
                        ET.SubElement(tp_node, 'icon', {'BUILTIN': 'stop'})
                    if tp_sts == 'b':
                        ET.SubElement(tp_node, 'icon', {'BUILTIN': 'prepare'})
                    if tp_sts == 'n':
                        ET.SubElement(tp_node, 'icon', {'BUILTIN': 'help'})
        fm_tree.write(fm_file.replace('.mm', '-TP.mm'))
        self.logger.info(self.log_prefix + \
                         "Successfully linked the test plan and execution results to file (%s)." % \
                         (fm_file.replace('.mm', '-TP.mm')))

    def _get_test_plan_info(self, name_filter, tc_list):
        self.logger.info(self.log_prefix + \
                         "Getting test plan and execution status from TestLink. This is going to take a while. Please wait...")
        self.tls = testlink.TestLinkHelper().connect(testlink.TestlinkAPIClient)
        prj = self.tls.getTestProjectByName(self.repo_name)
        prj_id = prj['id']
        tp_list = self.tls.getProjectTestPlans(prj_id)
        self.logger.info(self.log_prefix + \
                         "There are totally %d test plan for this project (%s)." % \
                         (len(tp_list), self.repo_name))
        for tp in tp_list:
            tp_name = tp['name']
            #TODO: Apply the name filter
            tp_id = tp['id']
            tc_dict = self.tls.getTestCasesForTestPlan(tp_id)
            for k in tc_dict.keys():
                tc = tc_dict[k][0]
                tc_id = tc['full_external_id']
                tc_sts = tc['exec_status']
                self._add_tc_history_list(tc_id, tc_sts, tp_name, tc_list)

        return 0

    def _add_tc_history_list(self, tc_id, tc_sts, tp_name, tc_list):
        for tc in tc_list:
            if tc_id == tc[0]:
                tc[1].append([tp_name, tc_sts])
                return True
        tc_list.append([tc_id, [[tp_name, tc_sts]]])
        return True

    def _remove_duplicate(self, old_list, new_list):
        for i in old_list:
            if not i in new_list:
                new_list.append(i)

    def _get_fm_tc_list(self, root_node, tc_list):
        for child in root_node.iter('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                tc_list.append(tc_id)

    def _update_fm_tp(self, root_node, tc_list):
        for child in root_node.findall('node'):
            for hook_node in child.findall('hook'):
                if hook_node.attrib['NAME'].strip() == 'accessories/plugins/AutomaticLayout.properties':
                    child.remove(hook_node)

            node_text = child.attrib['TEXT'].strip()
            #print node_text
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                if tc_id in tc_list:
                    child.attrib['COLOR'] = '#000000'
                else:
                    child.attrib['COLOR'] = '#cccccc'

            if not self._has_valid_tc_node(child, tc_list):
                child.attrib['COLOR'] = '#cccccc'
                child.attrib['FOLDED'] = 'true'
            else:
                child.attrib['COLOR'] = '#000000'
                child.attrib['FOLDED'] = 'false'
            self._update_fm_tp(child, tc_list)

        return 0

    def _has_valid_tc_node(self, root_node, tc_list):
        for child in root_node.iter('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if (tc_id.count(self.repo_prefix) == 1) and (tc_id in tc_list):
                return True
        return False

    def _remove_node_wo_tc(self, root_node):
        for child in root_node.findall('node'):
            if not self._has_tc_node(child):
                root_node.remove(child)
            else:
                self._remove_node_wo_tc(child)

        return 0

    def _has_tc_node(self, root_node):
        for child in root_node.iter('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                return True
        return False

    def _get_tc_list(self, root_node, exclude_tc_list, kept_tc_list, tc_list, ver_team, regression_level='5'):
        for child in root_node.findall('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            node_reg_lvl = regression_level
            for node_icon in child.findall('icon'):
                if node_icon.attrib['BUILTIN'].strip().count('full-') == 1:
                    node_reg_lvl = node_icon.attrib['BUILTIN'].strip()[-1]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                # TODO: If we want to implement verification team, we need add this information in this node  
                # Keep the node if regression level is matched and not in the exclude_tc_list, or it's in the must keep list kept_tc_list    
                if ((tc_id not in exclude_tc_list) and (int(node_reg_lvl) <= int(regression_level))) \
                        or (tc_id in kept_tc_list):
                    tc_list.append(tc_id)
                    #child.attrib['COLOR'] = '#000000'          
                    #else:
                    #child.attrib['COLOR'] = '#cccccc'
            else:
                self._get_tc_list(child, exclude_tc_list, kept_tc_list, tc_list, ver_team, node_reg_lvl)

        return 0

    def _update_tp(self, root_node, ver_team, exclude_tc_list, regression_level='5'):
        for child in root_node.findall('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            node_reg_lvl = regression_level
            for reg_lvl_icon in child.findall('icon'):
                if reg_lvl_icon.attrib['BUILTIN'].strip().count('full-') == 1:
                    node_reg_lvl = reg_lvl_icon.attrib['BUILTIN'].strip()[-1]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                # TODO: If we want to implement verification team, we need add this information in this node
                #tc_ver_team = node_text.split(PREFIX_TITLE_SEP)[1].split('|')
                if (tc_id in exclude_tc_list) or (int(node_reg_lvl) > int(regression_level)):
                    root_node.remove(child)
                    print node_text
            else:
                self._update_tp(child, ver_team, exclude_tc_list, node_reg_lvl)

        return 0

    def _find_removed_kept_tc(self, root_node, removed_tc_list, kept_tc_list):
        for child in root_node.findall('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                for icon_node in child.findall('icon'):
                    # TODO: How about multiple icons?
                    if icon_node.attrib['BUILTIN'].strip() == 'button_cancel':
                        removed_tc_list.append(tc_id)
                    if icon_node.attrib['BUILTIN'].strip() == 'button_ok':
                        kept_tc_list.append(tc_id)
            else:
                self._find_removed_kept_tc(child, removed_tc_list, kept_tc_list)

    def _find_removed_tc(self, root_node, tp_root, removed_tc_list):
        for child in root_node.findall('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                parent_text = root_node.attrib['TEXT'].strip()
                self.node_found = False
                self._find_combined_node(parent_text, node_text, tp_root)
                if not self.node_found:
                    removed_tc_list.append(tc_id)
            else:
                self._find_removed_tc(child, tp_root, removed_tc_list)

    def _find_combined_node(self, parent_text, tc_text, root_node):
        for child in root_node.findall('node'):
            node_text = child.attrib['TEXT'].strip()
            tc_id = node_text.split(PREFIX_TITLE_SEP)[0]
            # If this is the node for a test case
            if tc_id.count(self.repo_prefix) == 1:
                # Check the node text to see if it matches thus we know if this node exist in new test plan
                if (root_node.attrib['TEXT'].strip() == parent_text) and (node_text == tc_text):
                    # Need a global variable for this, haven't find a new way to replace this.
                    self.node_found = True
            else:
                self._find_combined_node(parent_text, tc_text, child)
        return 0

    def _get_link_node(self, node, link_list):
        for child in node.findall('node'):
            if child.attrib.has_key('LINK'):
                tds_id = node.attrib['TEXT'].split(' ')[0]
                tc_id = child.attrib['TEXT'].split(':')[0]
                tc_title = ''.join(child.attrib['TEXT'].split(':')[1:])
                link_list.append([tc_id, tc_title, tds_id])
            else:
                self._get_link_node(child, link_list)
        return 0

    def _remove_node_prefix(self, node):
        for child in node.iter('node'):
            # Make sure this is not the test case or requirement link node since only they are nodes with links
            if child.attrib.has_key('LINK') and child.attrib['LINK'].startswith(self.testlink_url):
                continue
            # If the node text is started with a number, then we consider it having added prefix
            # if child.attrib['TEXT'][0].isdigit:
            #     # Since Unicode may also be considered as numbers, we need to make sure this is unicode or prefix
            #     if (child.attrib['TEXT'].split(PREFIX_TITLE_SEP)[0] <> child.attrib['TEXT']):
            if not child.attrib.has_key('TEXT'):
                self.logger.error(self.log_prefix + \
                                 "Please check node (%s) since it may use a long name. Please convert it to plain text via FreeMind Menu Format=>Use Plaine Text." % \
                                 (child.attrib['ID'].strip()))
                exit(-1)
            if child.attrib['TEXT'].count(PREFIX_TITLE_SEP) == 0:
                continue
            self.logger.debug(self.log_prefix + \
                              "Prefix of node (%s) has been removed" % \
                              (child.attrib['TEXT']))
            child.attrib['TEXT'] = ''.join(child.attrib['TEXT'].split(PREFIX_TITLE_SEP)[1:])

        return 0

    def _remove_link_node(self, node):
        '''The key here is to use findall method since it will create a new children list'''
        for child in node.findall('node'):
            if child.attrib.has_key('LINK'):
                self.logger.debug(self.log_prefix + \
                                  "Link node (%s) has been removed from parent node (%s)" % \
                                  (child.attrib['TEXT'], node.attrib['TEXT']))
                node.remove(child)

            else:
                self._remove_link_node(child)

    def _add_node_prefix(self, node, num):
        ''' Add the node prefix (something like 1.1.2.1) for the TDS document
        '''
        res = 0
        i = 0
        for child in node.findall('node'):
            if child.attrib.has_key('LINK') and child.attrib['LINK'].startswith(self.testlink_url):
                continue
            i += 1
            prefix = str(num) + '.' + str(i)
            if child.attrib['TEXT'].count(PREFIX_TITLE_SEP) > 0:
                child.attrib['TEXT'] = prefix[4:] + PREFIX_TITLE_SEP + child.attrib['TEXT'].split(PREFIX_TITLE_SEP)[1:]
            else:
                child.attrib['TEXT'] = prefix[4:] + PREFIX_TITLE_SEP + child.attrib['TEXT']
            res = self._add_node_prefix(child, prefix)

        return res

    def extract_tc_from_docx(self, file_name, review_info):
        self.logger.info(self.log_prefix + \
                         "Reading test cases from file (%s). This is going to take a while. Please wait..." % \
                         file_name)

        if os.path.splitext(file_name)[-1] != '.docx':
            self.logger.error(self.log_prefix + \
                              "I am sorry that I can not parse this file. Please convert it to a docx file.")
            exit(-1)

        review_info = review_info.split('|')
        review_info = [item.strip() for item in review_info]
        if review_info == ['']:
            review_info = ['', '', '']

        tc_root = lxmlET.Element('testsuite', {'name': ''})
        lxmlET.SubElement(tc_root, 'node_order').text = lxmlET.CDATA('')
        lxmlET.SubElement(tc_root, 'details').text = lxmlET.CDATA('')

        ts_name = os.path.split(os.path.splitext(file_name)[0])[-1]
        child_ts_node = lxmlET.SubElement(tc_root, 'testsuite', {'name': ts_name})
        lxmlET.SubElement(child_ts_node, 'node_order').text = lxmlET.CDATA('')
        lxmlET.SubElement(child_ts_node, 'details').text = lxmlET.CDATA('')

        document = Document(file_name)
        tc_node_order = -1
        for table in document.tables:
            if table.cell(0, 0).paragraphs[0].text != 'Test case ID':
                continue
            tc_node_order += 1
            col_index = len(table.columns) - 2
            tc_id = table.cell(0, col_index).paragraphs[0].text.strip()
            print tc_id
            tc_purpose = '\n'.join([paragraph.text.strip() for paragraph in table.cell(1, col_index).paragraphs])
            tc_cfg = 'Test Configuration：\n'+ \
                     '\n'.join([paragraph.text.strip() for paragraph in table.cell(2, col_index).paragraphs])
            tc_pre_cond = 'Precondition：\n'+ \
                          '\n'.join([paragraph.text.strip() for paragraph in table.cell(3, col_index).paragraphs])
            tc_post_cond = 'Postcondition：\n'+ \
                           '\n'.join([paragraph.text.strip() for paragraph in table.cell(4, col_index).paragraphs])
            #print '\n'.join([tc_id, tc_purpose, tc_cfg, tc_pre_cond, tc_post_cond])
            testcase = lxmlET.SubElement(child_ts_node, 'testcase', {'name': tc_id})
            lxmlET.SubElement(testcase, 'node_order').text = lxmlET.CDATA(str(tc_node_order))
            lxmlET.SubElement(testcase, 'externalid').text = lxmlET.CDATA('')
            lxmlET.SubElement(testcase, 'version').text = lxmlET.CDATA('1')
            lxmlET.SubElement(testcase, 'summary').text = lxmlET.CDATA(self._replace_new_line(tc_purpose))
            lxmlET.SubElement(testcase, 'preconditions').text = lxmlET.CDATA(self._replace_new_line(tc_cfg + '\n' +\
                                                            tc_pre_cond + '\n' + tc_post_cond))
            lxmlET.SubElement(testcase, 'execution_type').text = lxmlET.CDATA('1')
            lxmlET.SubElement(testcase, 'importance').text = lxmlET.CDATA('3')

            steps = lxmlET.SubElement(testcase, 'steps')
            for i in range(6, len(table.rows)):
                step = lxmlET.SubElement(steps, 'step')
                lxmlET.SubElement(step, 'step_number').text = lxmlET.CDATA(str(i-5))
                action = '\n'.join([paragraph.text.strip() for paragraph in table.cell(i, 0).paragraphs])
                lxmlET.SubElement(step, 'actions').text = lxmlET.CDATA(self._replace_new_line(action))
                result = '\n'.join([paragraph.text.strip() for paragraph in table.cell(i, 1).paragraphs])
                lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA(self._replace_new_line(result))
                lxmlET.SubElement(step, 'execution_type').text = lxmlET.CDATA('1')

            custom_fields = lxmlET.SubElement(testcase, 'custom_fields')
            custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
            lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Regression Level')
            lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('5 - First Time Run|4 - Full Regression|3 - Regular Regression')
            custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
            lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Test Team')
            lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA('SIT')
            custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
            lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed')
            lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[0])
            custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
            lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed Version')
            lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[1])
            custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
            lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Review Info')
            lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[2])

        output_file_name = file_name.replace(os.path.splitext(file_name)[-1], '.xml')
        f = open(output_file_name, 'w')
        f.write(lxmlET.tostring(tc_root, xml_declaration=True, encoding='UTF-8', pretty_print=True))
        f.close
        self.logger.info(self.log_prefix + \
                         "Successfully generated test case file (%s). You can now import it into TestLink" % \
                         (output_file_name))

    def extract_tc_from_file(self, file_name, sheet_name, review_info):
        if os.path.splitext(file_name)[-1] == '.xls':
            self.extract_tc_from_xls(file_name, sheet_name, review_info)
        elif os.path.splitext(file_name)[-1].count('.doc') > 0:
            self.extract_tc_from_docx(file_name, review_info)

    def extract_tc_from_xls(self, file_name, sheet_name, review_info):
        xls_col_dict = {'TS_Name': -1, 'TS_Details': -1, 'Name': -1, 'Summary': -1, 'Preconditions': -1,
                        'Test Execution Type': -1, 'Importance': -1, 'HGI Regression Level': -1,
                        'HGI Test Team': -1, 'Steps': -1, 'Expected Results': -1, 'Step Execution Type': -1
                        , 'Requirements': -1}
        execution_type_dict = {'Manual': '1', 'Automated': '2'}
        importance_dict = {'H': '3', 'M': '2', 'L': '1'}
        regression_level_list = '5 - First Time Run|4 - Full Regression|3 - Regular Regression|2 - Basic Regression|1 - Basic Sanity'.split('|')

        if not os.path.exists(file_name):
            self.logger.error(self.log_prefix + \
                              "Cannot find the specified file (%s). Action aborted." % \
                              (file_name))
            return None
        self.logger.info(self.log_prefix + \
                         "Reading test cases from file (%s). This is going to take a while. Please wait..." % \
                         (file_name))
        src_wb = open_workbook(file_name, on_demand=True)

        sheet_name = sheet_name.split('|')
        sheet_name = [item.strip() for item in sheet_name]
        review_info = review_info.split('|')
        review_info = [item.strip() for item in review_info]
        if review_info == ['']:
            review_info = ['', '', '']
        for s in src_wb.sheets():
            if sheet_name <> [''] and s.name not in sheet_name:
                continue
            tc_root = lxmlET.Element('testsuite', {'name': ''})
            lxmlET.SubElement(tc_root, 'node_order').text = lxmlET.CDATA('')
            lxmlET.SubElement(tc_root, 'details').text = lxmlET.CDATA('')

            src_sheet = src_wb.sheet_by_name(s.name)
            ts_node = lxmlET.SubElement(tc_root, 'testsuite', {'name': s.name})
            child_ts_node = ts_node
            lxmlET.SubElement(ts_node, 'node_order').text = lxmlET.CDATA('')
            lxmlET.SubElement(ts_node, 'details').text = lxmlET.CDATA('')

            for i, cell in enumerate(src_sheet.col(0)):
                if i < 1:
                    continue
                if i == 1:
                    # Update the column index
                    for j in range(0, src_sheet.ncols):
                        cell_value = src_sheet.cell_value(i, j).strip()
                        if xls_col_dict.has_key(cell_value):
                            xls_col_dict[cell_value] = j
                    continue

                ts_name = src_sheet.cell_value(i, xls_col_dict['TS_Name']).strip()
                if ts_name <> '':
                    child_ts_node = lxmlET.SubElement(ts_node, 'testsuite', {'name': ts_name})
                    lxmlET.SubElement(child_ts_node, 'node_order').text = lxmlET.CDATA('')
                    lxmlET.SubElement(child_ts_node, 'details').text = lxmlET.CDATA(
                        self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['TS_Details']).strip()))
                tc_name = src_sheet.cell_value(i, xls_col_dict['Name']).strip()
                if tc_name <> '':
                    step_number = 1
                    testcase = lxmlET.SubElement(child_ts_node, 'testcase', {'name': tc_name})
                    lxmlET.SubElement(testcase, 'node_order').text = lxmlET.CDATA('')
                    lxmlET.SubElement(testcase, 'externalid').text = lxmlET.CDATA('')
                    lxmlET.SubElement(testcase, 'version').text = lxmlET.CDATA('1')
                    lxmlET.SubElement(testcase, 'summary').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Summary']).strip()))
                    lxmlET.SubElement(testcase, 'preconditions').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Preconditions']).strip()))
                    if not execution_type_dict.has_key(src_sheet.cell_value(i, xls_col_dict['Test Execution Type']).strip()):
                        self.logger.error(self.log_prefix + \
                                         "Wrong test case execution type (%s) in row(%d), col(%d) in sheet(%s) of file(%s)" % \
                                         (src_sheet.cell_value(i, xls_col_dict['Test Execution Type']).strip(), i+1, xls_col_dict['Test Execution Type']+1, s.name, file_name))
                        return
                    lxmlET.SubElement(testcase, 'execution_type').text = lxmlET.CDATA(execution_type_dict[src_sheet.cell_value(i, xls_col_dict['Test Execution Type']).strip()])
                    if not importance_dict.has_key(src_sheet.cell_value(i, xls_col_dict['Importance']).strip()):
                        self.logger.error(self.log_prefix + \
                                         "Wrong importance type (%s) in row(%d), col(%d) in sheet(%s) of file(%s)" % \
                                         (src_sheet.cell_value(i, xls_col_dict['Importance']).strip(), i+1, xls_col_dict['Importance']+1, s.name, file_name))
                        return
                    lxmlET.SubElement(testcase, 'importance').text = lxmlET.CDATA(importance_dict[src_sheet.cell_value(i, xls_col_dict['Importance']).strip()])
                    #lxmlET.SubElement(testcase, 'status').text = lxmlET.CDATA('Final')

                    steps = lxmlET.SubElement(testcase, 'steps')
                    step = lxmlET.SubElement(steps, 'step')
                    lxmlET.SubElement(step, 'step_number').text = lxmlET.CDATA(str(step_number))
                    lxmlET.SubElement(step, 'actions').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Steps']).strip()))
                    lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Expected Results']).strip()))
                    if not execution_type_dict.has_key(src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip()):
                        self.logger.error(self.log_prefix + \
                                         "Wrong test step execution type (%s) in row(%d), col(%d) in sheet(%s) of file(%s)" % \
                                         (src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip(), i+1, xls_col_dict['Step Execution Type']+1, s.name, file_name))
                        return
                    lxmlET.SubElement(step, 'execution_type').text = lxmlET.CDATA(execution_type_dict[src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip()])

                    custom_fields = lxmlET.SubElement(testcase, 'custom_fields')
                    custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
                    lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Regression Level')
                    regression_level = int(src_sheet.cell_value(i, xls_col_dict['HGI Regression Level']))
                    regression_level = '|'.join(regression_level_list[:len(regression_level_list) - regression_level + 1])
                    lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(regression_level)
                    custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
                    lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('HGI Test Team')
                    lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(src_sheet.cell_value(i, xls_col_dict['HGI Test Team']).strip())
                    custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
                    lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed')
                    lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[0])
                    custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
                    lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Reviewed Version')
                    lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[1])
                    custom_field = lxmlET.SubElement(custom_fields, 'custom_field')
                    lxmlET.SubElement(custom_field, 'name').text = lxmlET.CDATA('Review Info')
                    lxmlET.SubElement(custom_field, 'value').text = lxmlET.CDATA(review_info[2])

                step_info = src_sheet.cell_value(i, xls_col_dict['Steps'])
                if step_info <> "":
                    step_number += 1
                    step = lxmlET.SubElement(steps, 'step')
                    lxmlET.SubElement(step, 'step_number').text = lxmlET.CDATA(str(step_number))
                    lxmlET.SubElement(step, 'actions').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Steps']).strip()))
                    lxmlET.SubElement(step, 'expectedresults').text = lxmlET.CDATA(self._replace_new_line(src_sheet.cell_value(i, xls_col_dict['Expected Results']).strip()))
                    if not execution_type_dict.has_key(src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip()):
                        self.logger.error(self.log_prefix + \
                                         "Wrong test step execution type (%s) in row(%d), col(%d) in sheet(%s) of file(%s)" % \
                                         (src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip(), i+1, xls_col_dict['Step Execution Type'] + 1, s.name, file_name))
                        return
                    lxmlET.SubElement(step, 'execution_type').text = lxmlET.CDATA(execution_type_dict[src_sheet.cell_value(i, xls_col_dict['Step Execution Type']).strip()])
                    # requirements = lxmlET.SubElement(testcase, 'requirements')
                    # requirement = lxmlET.SubElement(requirements, 'requirement')
                    # lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
                    #     os.path.splitext(os.path.split(self.tds_url)[-1])[0])
                    # lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(tc_tds_dict[tds_item.attrib['ID']][0])
                    # if not tc_pfs_dict.has_key(tds_item.attrib['ID']):
                    #     return
                    # for pfs_id in tc_pfs_dict[tds_item.attrib['ID']]:
                    #     requirement = lxmlET.SubElement(requirements, 'requirement')
                    #     lxmlET.SubElement(requirement, 'req_spec_title').text = lxmlET.CDATA(
                    #         os.path.splitext(os.path.split(self.pfs_url)[-1])[0])
                    #     lxmlET.SubElement(requirement, 'doc_id').text = lxmlET.CDATA(pfs_id)

            output_file_name = file_name.replace(os.path.splitext(file_name)[-1], '_' + s.name + '.xml')
            f = open(output_file_name, 'w')
            f.write(lxmlET.tostring(tc_root, xml_declaration=True, encoding='UTF-8', pretty_print=True))
            f.close
            self.logger.info(self.log_prefix + \
                             "Successfully generated test case file (%s). You can now import it into TestLink" % \
                             (output_file_name))

    def _replace_new_line(self, text):
        return '<p>' + text.replace('\n', '</p><p>') + '</p>'

    def extract_requirements(self, req_file_name, template):
        pmr_list = []
        pfs_list = []
        pfs_pmr_list = []
        pmr_pfs_list = []
        prefixed_pmr_pfs_list = []

        if not os.path.exists(req_file_name):
            self.logger.error(self.log_prefix + \
                              "Cannot find the specified file (%s). Action aborted." % \
                              (req_file_name))
            return None

        if template == 'KreaTV':
            res = self._read_req_from_xls_kreatv(req_file_name, pmr_list, pfs_list, pfs_pmr_list)
        else:
            if os.path.splitext(req_file_name)[-1] in ['.doc', '.docx']:
                res = self._read_req_from_docx_hgi(req_file_name, pmr_list, pfs_list, pfs_pmr_list)
            else:
                res = self._read_req_from_xls_hgi(req_file_name, pmr_list, pfs_list, pfs_pmr_list)

        if len(pfs_pmr_list) > 0:
            res = self._reverse_links(pfs_pmr_list, pmr_pfs_list)
            res = self._add_req_prefix(pmr_pfs_list, prefixed_pmr_pfs_list)

        # Get the filename without extension.
        title = os.path.splitext(os.path.split(self.pmr_url)[-1])[0]
        if len(pmr_list) > 0:
            res = self._gen_req_xml(pmr_list, title, self.pmr_url, self.pmr_prefix, prefixed_pmr_pfs_list)
        title = os.path.splitext(os.path.split(self.pfs_url)[-1])[0]
        res = self._gen_req_xml(pfs_list, title, self.pfs_url, self.pfs_prefix, prefixed_pmr_pfs_list)

        title = os.path.splitext(os.path.split(self.pmr_url)[-1])[0]
        if len(pmr_list) > 0:
            res = self._gen_req_freemind(pmr_list, title, self.pmr_url.replace('.xml', '.mm'), self.pmr_prefix)
        title = os.path.splitext(os.path.split(self.pfs_url)[-1])[0]
        res = self._gen_req_freemind(pfs_list, title, self.pfs_url.replace('.xml', '.mm'), self.pfs_prefix)

        if len(pfs_pmr_list) > 0:
            res = self._build_fm_traceability(self.pfs_url.replace('.xml', '.mm'), self.pmr_url.replace('.xml', '.mm'),
                                              pfs_pmr_list, self.pfs_url.replace('.xml', '[PFS-PMR].mm'))
            res = self._build_fm_traceability(self.pmr_url.replace('.xml', '.mm'), self.pfs_url.replace('.xml', '.mm'),
                                              pmr_pfs_list, self.pmr_url.replace('.xml', '[PMR-PFS].mm'))
        return res

    def _add_req_prefix(self, pmr_pfs_list, prefixed_pmr_pfs_list):
        for i, pmr_item in enumerate(pmr_pfs_list):
            prefixed_pmr_pfs_list.append([self.pmr_prefix + pmr_item[0], []])
            for pfs_item in pmr_item[1]:
                prefixed_pmr_pfs_list[i][1].append(self.pfs_prefix + pfs_item)
        return 0

    def _reverse_links(self, orig_list, reversed_list):
        ''' The original list is something like [PFS_ID, [PMR_ID1, PMRID2,...]].
            The reversed list is something like [PMR_ID, [PFS_ID1, PFS_ID2]]
        '''
        self.logger.debug(self.log_prefix + \
                          "Reversing the traceability links.")
        for orig_link in orig_list:
            src_id = orig_link[0]
            for link_id in orig_link[1]:
                if link_id == '':
                    continue
                reversed_link_exist = False
                i = 0
                for i, reversed_link in enumerate(reversed_list):
                    if reversed_link[0] == link_id:
                        reversed_link_exist = True
                        break
                if reversed_link_exist:
                    reversed_list[i][1].append(src_id)
                else:
                    reversed_list.append([link_id, [src_id]])

                    #pprint.pprint(pmr_pfs_list)
        return 0

    def _build_fm_traceability(self, dst_fm, src_fm, link_list, output_file, tds_file=False):
        ''' This function is using to two FreeMind maps by using the traceability list in link_list[]
            link_list[] has the format of either [PFS_ID, [PMR_ID1, PMRID2,...]] or [PMR_ID, [PFS_ID1, PFS_ID2]] depends on 
            what's the destination FreeMind map.
        '''
        self.logger.info(self.log_prefix + \
                         "Building the FreeMind traceability file %s (Between %s and %s)." % \
                         (output_file, dst_fm, src_fm))
        dst_fm_tree = ET.parse(dst_fm)
        dst_fm_root = dst_fm_tree.getroot()
        src_fm_root = ET.parse(src_fm).getroot()
        new_added_nodes = []

        for dst_node in dst_fm_root.iter('node'):
            if tds_file:
                if not self._last_tds_node(dst_node):
                    continue
            else:
                if dst_node.find('node') is not None:
                    continue
            # Please note the new added nodes will be looped through iter again so we need to ignore that by using new_added_nodes[]
            if dst_node.attrib['TEXT'] not in new_added_nodes:
                if tds_file:
                    dst_id = dst_node.attrib['ID'].strip()
                else:
                    dst_id = dst_node.attrib['TEXT'].strip().split(PREFIX_TITLE_SEP)[0]
                traceability_links = []
                for traceability in link_list:
                    if dst_id == traceability[0]:
                        traceability_links = traceability[1]
                        break
                if (traceability_links == []) or (traceability_links == ['']):
                    # Highlight the node with traceability missing
                    self.logger.warning(self.log_prefix + \
                                        "Highlight the node (%s) with missing traceability for file %s." % \
                                        (dst_node.attrib['TEXT'].strip(), output_file))
                    dst_node.set('BACKGROUND_COLOR', '#ff0000')
                for link_id in traceability_links:
                    if link_id == '':
                        continue
                    link_found = False
                    for src_node in src_fm_root.iter('node'):
                        if (src_node.attrib['TEXT'].split(PREFIX_TITLE_SEP)[0] == link_id):
                            link_found = True
                            dst_node.append(src_node)
                            new_added_nodes.append(src_node.attrib['TEXT'])
                            self.logger.debug(self.log_prefix + \
                                              "Add link %s to %s." % \
                                              (link_id, dst_id))
                            break
                    if not link_found:
                        self.logger.warning(self.log_prefix + \
                                            "Cannot find link %s for %s for file %s." % \
                                            (link_id, dst_id, output_file))
                        # Highlight the node with traceability missing
                        self.logger.warning(self.log_prefix + \
                                            "Highlight the node (%s) with missing traceability for file %s." % \
                                            (dst_node.attrib['TEXT'].strip(), output_file))
                        dst_node.set('BACKGROUND_COLOR', '#ff0000')

        dst_fm_tree.write(output_file)

        self.logger.info(self.log_prefix + \
                         "Successfully built the FreeMind traceability file %s (Between %s and %s)." % \
                         (output_file, dst_fm, src_fm))

        return 0

    def _link_pfs_pmr(self, dst_fm, src_fm, link_list, output_file):
        ''' This function is using to link PMR FreeMind map and PFS FreeMind map by using the traceability list in link_list[]
            link_list[] has the format of either [PFS_ID, [PMR_ID1, PMRID2,...]] or [PMR_ID, [PFS_ID1, PFS_ID2]] depends on 
            what's the destination FreeMind map.
        '''
        dst_fm_tree = ET.parse(dst_fm)
        dst_fm_root = dst_fm_tree.getroot()
        src_fm_root = ET.parse(src_fm).getroot()
        new_added_nodes = []

        for dst_node in dst_fm_root.iter('node'):
            # Please note the new added nodes will be looped through iter again so we need to ignore that by using new_added_nodes[]
            if dst_node.attrib.has_key('LINK') and (dst_node.attrib['TEXT'] not in new_added_nodes):
                req_id = dst_node.attrib['TEXT'].split(PREFIX_TITLE_SEP)[0]
                req_links = []
                for req_trace in link_list:
                    if req_id == req_trace[0]:
                        req_links = req_trace[1]
                        break
                if (req_links == []) or (req_links == ['']):
                    # Highlight the node with traceability missing
                    self.logger.warning(self.log_prefix + \
                                        "Cannot find the requirement links for %s." % \
                                        (req_id))
                    dst_node.set('BACKGROUND_COLOR', '#ff0000')
                for req_link_id in req_links:
                    if req_link_id == '':
                        continue
                    link_found = False
                    for src_node in src_fm_root.iter('node'):
                        if src_node.attrib.has_key('LINK') and (
                                    src_node.attrib['TEXT'].split(PREFIX_TITLE_SEP)[0] == req_link_id):
                            link_found = True
                            dst_node.append(src_node)
                            new_added_nodes.append(src_node.attrib['TEXT'])
                            self.logger.info(self.log_prefix + \
                                             "Add requirement link %s to %s." % \
                                             (req_link_id, req_id))
                            break
                    if not link_found:
                        self.logger.error(self.log_prefix + \
                                          "Cannot find requirement link %s for %s." % \
                                          (req_link_id, req_id))

        dst_fm_tree.write(output_file)

        return 0

    def _gen_req_freemind(self, req_list, title, output_file, prefix):
        ''' req_list is a list like [GROUP_NAME, [ [REQ_ID, REQ_TITLE, REQ_DESC, REQ_VER_TEAM], ... ] ]
            REQ_ID and REQ_TITLE will be combined as the node text and REQ_DESC will be displayed as comments
        '''
        self.logger.info(self.log_prefix + \
                         "Generating the FreeMind file %s (Document Title: %s. Document ID Prefix: %s)." % \
                         (output_file, title, prefix))
        freemind = lxmlET.Element('map', {'version': '1.0.1'})

        lxmlET.SubElement(freemind, 'attribute_registry', {'SHOW_ATTRIBUTES': 'hide'})
        root_node = lxmlET.SubElement(freemind, 'node',
                                      {'BACKGROUND_COLOR': '#0000ff', 'COLOR': '#000000', 'TEXT': title})
        lxmlET.SubElement(root_node, 'font', {'NAME': 'SansSerif', 'SIZE': '20'})
        lxmlET.SubElement(root_node, 'hook', {'NAME': 'accessories/plugins/AutomaticLayout.properties'})

        req_count = 0
        for group in req_list:
            group_node = lxmlET.SubElement(root_node, 'node', {'COLOR': '#990000', 'FOLDED': "true", 'TEXT': group[0]})
            i = 0
            for i, req_item in enumerate(group[1]):
                node_text = req_item[REQ_ID] + PREFIX_TITLE_SEP + req_item[REQ_VER_TEAM] + PREFIX_TITLE_SEP + req_item[
                    REQ_TITLE]
                node_comment = req_item[REQ_DESC]
                node_link = self.testlink_url + '/linkto.php?tprojectPrefix=' + self.repo_prefix + '&item=req&id=' + prefix + \
                            req_item[REQ_ID]
                req_node = lxmlET.SubElement(group_node, 'node',
                                             {'COLOR': '#990000', 'LINK': node_link, 'TEXT': node_text})
                richcontent = lxmlET.SubElement(req_node, 'richcontent', {'TYPE': 'NOTE'})
                html = lxmlET.SubElement(richcontent, 'html')
                lxmlET.SubElement(richcontent, 'head')
                body = lxmlET.SubElement(html, 'body')
                comment = lxmlET.SubElement(body, 'p')
                comment.text = node_comment
            i = i + 1
            req_count = req_count + i
            group_node.attrib['TEXT'] = group_node.attrib['TEXT'] + '[' + str(i) + ']'

        root_node.attrib['TEXT'] = root_node.attrib['TEXT'] + '[' + str(req_count) + ']'

        #self._update_pfs_node_format(freemind)
        lxmlET.ElementTree(freemind).write(output_file)
        self.logger.info(self.log_prefix + \
                         "Successfully generated the FreeMind file %s (Document Title: %s. Document ID Prefix: %s)." % \
                         (output_file, title, prefix))
        return 0

    def _read_req_from_docx_hgi(self, file_name, pmr_list, pfs_list, trace_list):
        """
        Read requirements from HGI SDS template
        :param file_name:
        :param pmr_list:
        :param pfs_list:
        :param trace_list:
        """
        self.logger.info(self.log_prefix + \
                         "Reading requirements from file (%s). This is going to take a while. Please wait..." % \
                         file_name)

        if os.path.splitext(file_name)[-1] != '.docx':
            self.logger.error(self.log_prefix + \
                              "I am sorry that I can not parse this file. Please convert it to a docx file.")
            exit(-1)
        pfs_index_list = []
        pfs_grp_list = []
        pfs_grp_id = 0
        valid_columns = ['Index', 'Category', 'Description', 'DEV', 'DVT', 'FT', 'SI&T', 'Comment']
        ver_team_list = ['DEV', 'DVT', 'FT', 'SIT']
        pfs_ver_team = ''
        document = Document(file_name)
        for table in document.tables:
            invalid_table = False
            if len(table.columns) != len(valid_columns):
                continue
            for i in range(0, len(table.rows)):
                pfs_item = []
                for j in range(0, len(table.columns)):
                    cell = table.cell(i, j)
                    paragraph_text = ''
                    for k, paragraph in enumerate(cell.paragraphs):
                        if i == 0 and paragraph.text != valid_columns[j]:
                            invalid_table = True
                            break
                        elif i > 0:
                            if k > 0:  #paragraph.style.startswith('List'):
                                paragraph_text += '\n'
                            paragraph_text += paragraph.text.strip()
                    pfs_item.append(paragraph_text)
                    if invalid_table:
                        break
                if invalid_table:
                    break
                if i == 0 or pfs_item[0] == '':
                    continue
                pfs_cat = pfs_item[1]
                if pfs_cat != '':
                    if pfs_cat in pfs_grp_list:
                        pfs_grp_id = pfs_grp_list.index(pfs_cat)
                    else:
                        pfs_grp_list.append(pfs_cat)
                        pfs_list.append([pfs_cat, []])
                        pfs_grp_id = len(pfs_grp_list) - 1
                if pfs_item[0] not in pfs_index_list:
                    pfs_ver_team = ''
                    for ver_index in range(0, len(ver_team_list)):
                        if pfs_item[3 + ver_index] == 'Y':
                            pfs_ver_team += '|' + ver_team_list[ver_index]
                    pfs_ver_team = '|'.join(pfs_ver_team.split('|')[1:])
                    pfs_phase = ''
                    if pfs_item[7].upper().startswith('P'):
                        pfs_phase = pfs_item[7]
                    pfs_list[pfs_grp_id][1].append(
                        [pfs_item[0], pfs_item[2], pfs_item[2], pfs_ver_team, '', pfs_phase])
                    pfs_index_list.append(pfs_item[0])
                else:
                    self.logger.error(self.log_prefix + "%s is duplicated." % pfs_item[0])
            if invalid_table:
                continue

        #pprint.pprint(pfs_list)
        self.logger.info(self.log_prefix + "%d PFS items and %d categories found in %s." % (
            len(pfs_index_list), len(pfs_grp_list), file_name))
        return 0

    def _read_req_from_xls_hgi(self, file_name, pmr_list, pfs_list, trace_list):
        """ This function will read a Excel and extract PMR, PFS and traceability out of it.
        """
        self.logger.info(self.log_prefix + \
                         "Reading requirements from file (%s). This is going to take a while. Please wait..." % \
                         file_name)

        if os.path.splitext(file_name)[-1] != '.xls':
            self.logger.error(self.log_prefix + \
                              "I am sorry that I can not parse this file. Please convert it to a xls file.")
            exit(-1)
        src_wb = open_workbook(file_name, on_demand=True, formatting_info=True)

        # The following columns are optional
        pfs_phase_col = -1
        pfs_ft_col = -1
        pmr_title_col = -1
        pfs_title_col = -1
        col_defined = False
        pmr_pfs_trace_list = []
        pmr_index_list = []
        pfs_index_list = []
        for s in src_wb.sheets():
            src_sheet = src_wb.sheet_by_name(s.name)
            if s.name.find('Specification') != -1:
                pmr_grp_id = 0
                pfs_grp_id = 0
                pre_pmr_index = ''

                for i, cell in enumerate(src_sheet.col(0)):
                    if not col_defined:
                        for j in range(0, src_sheet.ncols):
                            cell_text = str(src_sheet.cell_value(i, j)).strip()
                            if cell_text.lower() == 'pmr index':
                                pmr_index_col = j
                            if cell_text.lower() == 'pmr title':
                                pmr_title_col = j
                            if cell_text.lower() == 'pmr description':
                                pmr_desc_col = j
                            if cell_text.lower() == 'index':
                                pfs_index_col = j
                            if cell_text.lower() == 'pfs title':
                                pfs_title_col = j
                            if cell_text.lower() == 'category':
                                pfs_cat_col = j
                            if cell_text.lower() == 'phase':
                                pfs_phase_col = j
                            if cell_text.lower() == 'description':
                                pfs_desc_col = j
                            if cell_text.lower() == 'dev':
                                pfs_dev_col = j
                            if cell_text.lower() == 'dvt':
                                pfs_dvt_col = j
                            if cell_text.lower() == 'si&t':
                                pfs_sit_col = j
                                col_defined = True
                            if cell_text.lower() == 'ft':
                                pfs_ft_col = j
                            if cell_text.lower().endswith('comments'):
                                pmr_cmt_col = j
                    else:
                        pmr_index = src_sheet.cell_value(i, pmr_index_col).strip()
                        pmr_desc = src_sheet.cell_value(i, pmr_desc_col).strip()
                        pmr_ver_team = 'ATP'
                        if pmr_index != '' and pmr_desc == '':
                            # This is a PMR category
                            pmr_grp_desc = src_sheet.cell_value(i, pmr_index_col).strip()
                            pmr_list.append([pmr_grp_desc, []])
                            pmr_grp_id = len(pmr_list) - 1
                        if len(pmr_list) == 0:
                            pmr_list.append(['Default Category', []])

                        pfs_index = src_sheet.cell_value(i, pfs_index_col).strip()
                        pfs_desc = src_sheet.cell_value(i, pfs_desc_col).strip()
                        pfs_cat = src_sheet.cell_value(i, pfs_cat_col).strip()
                        pfs_dev = src_sheet.cell_value(i, pfs_dev_col).strip()
                        pfs_dvt = src_sheet.cell_value(i, pfs_dvt_col).strip()
                        pfs_sit = src_sheet.cell_value(i, pfs_sit_col).strip()
                        pmr_cmt = src_sheet.cell_value(i, pmr_cmt_col).strip()
                        if pmr_cmt != '':
                            pmr_cmt = 'SE Comments:' + pmr_cmt

                        pmr_title = ''
                        if pmr_title_col != -1:
                            pmr_title = src_sheet.cell_value(i, pmr_title_col).strip()
                        pfs_title = ''
                        if pfs_title_col != -1:
                            pfs_title = src_sheet.cell_value(i, pfs_title_col).strip()

                        pfs_ft = ''
                        if pfs_ft_col != -1:
                            # This is an optional column
                            pfs_ft = src_sheet.cell_value(i, pfs_ft_col).strip()
                        pfs_phase = ''
                        if pfs_phase_col != -1:
                            # This is an optional column
                            pfs_phase = str(src_sheet.cell_value(i, pfs_phase_col)).strip()
                            if not pfs_phase.upper().startswith('P'):
                                pfs_phase = 'P' + pfs_phase
                                pfs_phase = pfs_phase[:2]

                        for merged_range in src_sheet.merged_cells:
                            rlo, rhi, clo, chi = merged_range
                            if (i >= rlo) and (i < rhi) and (pfs_index_col >= clo) and (pfs_index_col < chi):
                                pfs_index = src_sheet.cell_value(rlo, pfs_index_col).strip()
                            if (i >= rlo) and (i < rhi) and (pfs_desc_col >= clo) and (pfs_desc_col < chi):
                                pfs_desc = src_sheet.cell_value(rlo, pfs_desc_col).strip()
                            if (i >= rlo) and (i < rhi) and (pmr_index_col >= clo) and (pmr_index_col < chi):
                                pmr_index = src_sheet.cell_value(rlo, pmr_index_col).strip()
                            if (i >= rlo) and (i < rhi) and (pmr_desc_col >= clo) and (pmr_desc_col < chi):
                                pmr_desc = src_sheet.cell_value(rlo, pmr_desc_col).strip()
                            if (i >= rlo) and (i < rhi) and (pfs_cat_col >= clo) and (pfs_cat_col < chi):
                                pfs_cat = src_sheet.cell_value(rlo, pfs_cat_col).strip()
                            if (i >= rlo) and (i < rhi) and (pmr_title_col >= clo) and (pmr_title_col < chi):
                                pmr_title = src_sheet.cell_value(rlo, pmr_title_col).strip()
                            if (i >= rlo) and (i < rhi) and (pfs_title_col >= clo) and (pfs_title_col < chi):
                                pfs_title = src_sheet.cell_value(rlo, pfs_title_col).strip()

                        if pmr_index == 'PMR Index':
                            continue

                        if pfs_cat != '':
                            pfs_cat_exist = False
                            for item_index, pfs_item in enumerate(pfs_list):
                                if pfs_cat == pfs_item[0]:
                                    # This is an existing PFS category
                                    pfs_grp_id = item_index
                                    pfs_cat_exist = True
                                    break
                            if not pfs_cat_exist:
                                # This is a new PFS category
                                pfs_list.append([pfs_cat, []])
                                pfs_grp_id = len(pfs_list) - 1

                        if pmr_title == '':
                            pmr_title = pmr_desc
                        if pfs_title == '':
                            pfs_title = pfs_desc

                        pfs_ver_team = ''
                        if pfs_dev.upper() == 'Y':
                            pfs_ver_team = '|DEV'
                        if pfs_dvt.upper() == 'Y':
                            pfs_ver_team += '|DVT'
                        if pfs_sit.upper() == 'Y':
                            pfs_ver_team += '|SIT'
                        if pfs_ft.upper() == 'Y':
                            pfs_ver_team += '|FT'
                        pfs_ver_team = '|'.join(pfs_ver_team.split('|')[1:])

                        if src_sheet.cell_value(i, pmr_index_col).strip() in pmr_index_list:
                            self.logger.error(self.log_prefix + \
                                              "%s on row %d is duplicated." % \
                                              (src_sheet.cell_value(i, pmr_index_col).strip(), i + 1))
                        if src_sheet.cell_value(i, pfs_index_col).strip() in pfs_index_list:
                            self.logger.error(self.log_prefix + \
                                              "%s on row %d is duplicated." % \
                                              ( src_sheet.cell_value(i, pfs_index_col).strip(), i + 1))

                        if pmr_index != '' and pmr_desc != '' and pfs_index != '':
                            # PFS item traced to PMR item
                            if pmr_index not in pmr_index_list:
                                pmr_list[pmr_grp_id][1].append(
                                    [pmr_index, pmr_title, pmr_desc, pmr_ver_team, pmr_cmt, ''])
                                pmr_index_list.append(pmr_index)
                            if pfs_index not in pfs_index_list:
                                pfs_list[pfs_grp_id][1].append(
                                    [pfs_index, pfs_title, pfs_desc, pfs_ver_team, '', pfs_phase])
                                pfs_index_list.append(pfs_index)
                            self._add_traceability(pmr_pfs_trace_list, pmr_index, [pfs_index])
                        if pmr_index == '' and pmr_desc == '' and pfs_index != '' and pfs_desc != '':
                            # New PFS item traced to previous PMR item
                            pmr_index = pre_pmr_index
                            if pfs_index not in pfs_index_list:
                                pfs_list[pfs_grp_id][1].append(
                                    [pfs_index, pfs_title, pfs_desc, pfs_ver_team, '', pfs_phase])
                                pfs_index_list.append(pfs_index)
                            if pre_pmr_index <> '':
                                self._add_traceability(pmr_pfs_trace_list, pmr_index, [pfs_index])
                        if pmr_index == '' and pmr_desc == '' and pfs_index == '' and pfs_desc != '':
                            # Traceability only PFS item traced to previous PMR item
                            pmr_index = pre_pmr_index
                            if pre_pmr_index <> '':
                                self._add_traceability(pmr_pfs_trace_list, pmr_index, pfs_desc.split('\n'))
                        if pmr_index != '' and pmr_desc != '' and pfs_index == '' and pfs_desc != '':
                            # Existing PFS item traced to new PMR item
                            if pmr_index not in pmr_index_list:
                                pmr_list[pmr_grp_id][1].append(
                                    [pmr_index, pmr_title, pmr_desc, pmr_ver_team, pmr_cmt, ''])
                                pmr_index_list.append(pmr_index)
                            self._add_traceability(pmr_pfs_trace_list, pmr_index, pfs_desc.split('\n'))
                        if pmr_index != '' and pmr_desc != '' and pfs_index == '' and pfs_desc == '':
                            # New PMR item with no PFS item
                            if pmr_index not in pmr_index_list:
                                pmr_list[pmr_grp_id][1].append(
                                    [pmr_index, pmr_title, pmr_desc, pmr_ver_team, pmr_cmt, ''])
                                pmr_index_list.append(pmr_index)
                        if pmr_index == '' and pmr_desc == '' and pfs_index != '' and pfs_desc != '':
                            # New PFS item without PMR item
                            if pfs_index not in pfs_index_list:
                                pfs_list[pfs_grp_id][1].append(
                                    [pfs_index, pfs_title, pfs_desc, pfs_ver_team, '', pfs_phase])
                                pfs_index_list.append(pfs_index)

                        if pmr_index != '':
                            pre_pmr_index = pmr_index
                        if pfs_index != '':
                            pre_pfs_index = pfs_index

        self._reverse_links(pmr_pfs_trace_list, trace_list)
        #pprint.pprint(pmr_list)
        self.logger.info(self.log_prefix + \
                         "Successfully extracted requirements from file (%s). %d PMR items and %d PFS items found." % \
                         (file_name, len(pmr_index_list), len(pfs_index_list)))
        return 0

    def _add_traceability(self, trace_list, dst_index, src_index_list):
        """
        This function is used to generate a traceablity list like [PMR, [PFS1, PFS2, PFS3]]
        :param trace_list:
        :param dst_index:
        :param src_index_list:
        """
        for trace_item in trace_list:
            if dst_index == trace_item[0]:
                for new_src_index in src_index_list:
                    duplicated_src = False
                    for orig_src_index in trace_item[1]:
                        if new_src_index == orig_src_index:
                            duplicated_src = True
                            break
                    if not duplicated_src:
                        trace_item[1].append(new_src_index)
                return
        trace_list.append([dst_index, src_index_list])

    def _read_req_from_xls_kreatv(self, file_name, pmr_list, pfs_list, trace_list):
        ''' This function will read a Excel and extract PMR, PFS and traceability out of it.
        '''
        self.logger.info(self.log_prefix + \
                         "Reading requirements from file (%s). This is going to take a while. Please wait..." % \
                         (file_name))
        src_wb = open_workbook(file_name, on_demand=True)

        for s in src_wb.sheets():
            src_sheet = src_wb.sheet_by_name(s.name)
            if s.name == 'PMR':
                group_id = 0
                for i, cell in enumerate(src_sheet.col(0)):
                    req_id = cell.value.strip()
                    req_title = src_sheet.cell_value(i, 1).strip()
                    req_desc = src_sheet.cell_value(i, 2).strip()
                    ver_team = 'ATP'
                    if req_desc == '':
                        group_id = group_id + 1
                        pmr_list.append([req_title, []])
                    else:
                        pmr_list[group_id - 1][1].append([req_id, req_title, req_desc, ver_team])
                        #pprint.pprint(pmr_list)
            if s.name == 'Requirements':
                group_id = 0
                for i, cell in enumerate(src_sheet.col(0)):
                    if i > 0:
                        req_id = cell.value.strip()
                        req_title = src_sheet.cell_value(i, 1).strip()
                        ver_team = src_sheet.cell_value(i, 3).strip()
                        req_desc = src_sheet.cell_value(i, 4).strip()
                        if req_desc == '':
                            group_id = group_id + 1
                            pfs_list.append([req_id, []])
                        else:
                            pfs_list[group_id - 1][1].append([req_id, req_title, req_desc, ver_team])
                            #pprint.pprint(pfs_list)
            if s.name == 'PFS':
                group_id = 0
                for i, cell in enumerate(src_sheet.col(0)):
                    if i > 0:
                        req_id = cell.value.strip()
                        req_trace = src_sheet.cell_value(i, 2).strip().split('\n')
                        if len(req_trace) == 1:
                            req_trace = src_sheet.cell_value(i, 2).strip().split(' ')
                        if len(req_trace) == 1:
                            req_trace = src_sheet.cell_value(i, 2).strip().split(',')
                        if len(req_trace) == 1:
                            req_trace = src_sheet.cell_value(i, 2).strip().split(';')
                            #req_trace = '|'.join(req_trace)
                        if str(src_sheet.cell_value(i, 1)).strip() <> '':
                            trace_list.append([req_id, req_trace])
                            #pprint.pprint(trace_list)
        self.logger.info(self.log_prefix + \
                         "Successfully extracted requirements from file (%s)." % \
                         (file_name))
        return 0


def args_parser(arguments=None):
    parser = argparse.ArgumentParser(description= \
                                         'This application can be used to extract event test case, sub-procedure test cases and\
        multi-procedure test cases information from a Excel file, and generate test cases based \
        on these configurations. It can also be used to parse C header files to get the struct\
        template and message map.')
    parser.add_argument('--version', action='version', version='%(prog)s 0.1')
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument('-ap', '--add_prefix', action='store_true',
                       help="Update the existing Freemind file and add the prefix to each node except for \
              nodes with links (for instance: test case node or requirement node). \
              The most common usage is FreeMind -ap -f FREEMIND_FILE.")
    group.add_argument('-rp', '--remove_prefix', action='store_true',
                       help="Update the existing Freemind file and remove the prefix to each node. \
              The most common usage is FreeMind -rp -f FREEMIND_FILE.")
    group.add_argument('-g', '--gen_tds', action='store_true',
                       help="Update the existing Freemind file and remove the prefix to each node. \
              The most common usage is FreeMind -rp -f FREEMIND_FILE.")
    group.add_argument('-l', '--link_tds', action='store_true',
                       help="Extract test case and TDS linkage information from xml file exported from TestLink.\
                and update the FreeMind file with test cases links.\
                The most common usage is FreeMind -l -f FREEMIND_FILE -xml XML_FILE.")

    parser.add_argument('-s', '--src_file',
                        help="Specify the FreeMind file which contains various nodes of test design specification.")

    parser.add_argument('-d', '--dst_file',
                        help="Specify the xml file exported from TestLink with test case and TDS linkage information.")

    if arguments == None:
        args = parser.parse_args()
    else:
        args = parser.parse_args(arguments)

    return args


def start_main():
    reload(sys)
    sys.setdefaultencoding('utf-8')
    logging.config.fileConfig(PKG_PATH + 'logging.conf')
    logger = logging.getLogger(__name__)
    cfg_file = './config.xml'
    if os.path.exists(cfg_file):
        FreeMind(logger, cfg_file)
        sys.exit()

    freemind = FreeMind(logger)
    args = args_parser()
    if (args.add_prefix and args.src_file != None):
        freemind.add_prefix(args.src_file)
        sys.exit()
    if (args.remove_prefix and args.src_file != None):
        freemind.remove_prefix(args.src_file)
        sys.exit()
    if (args.gen_tds and args.src_file != None):
        freemind.gen_tds(args.src_file)
        sys.exit()
    if (args.link_tds and args.src_file != None and args.dst_file != None):
        if os.path.splitext(args.src_file)[-1] == 'mm' and os.path.splitext(args.dst_file)[-1] == 'xml':
            freemind.link_tds2tc(args.src_file, args.dst_file)
        if os.path.splitext(args.src_file)[-1] == 'xml' and os.path.splitext(args.dst_file)[-1] == 'mm':
            freemind.link_tc2tds(args.dst_file, args.src_file)
        sys.exit()


if __name__ == '__main__':
    start_main()